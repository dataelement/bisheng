"""python-docx helpers for building .docx inside the BiSheng code interpreter.

Use from a build script:

    import sys
    sys.path.insert(0, "skills/bisheng-docx/scripts")
    from docx_helpers import set_run_font, apply_chinese_defaults, add_table

Every function here exists because plain python-docx gets the Chinese case
wrong or cannot express the construct at all:

* ``run.font.name`` writes only ``w:ascii``/``w:hAnsi`` — Chinese glyphs keep
  falling back to the theme font. The CJK face lives in ``w:eastAsia`` and has
  to be set through the XML.
* Table column widths must be written to **every cell**; setting them on the
  column object alone is silently ignored by Word.
* A table of contents, page numbers, and a horizontal rule are all field codes
  or borders with no python-docx API at all.

**Which face and size a helper renders in lives in ``style_profiles.py``, not
here.** Call ``apply_chinese_defaults(doc)`` once — it selects the GB/T 9704
公文 profile by default — and every helper below follows it. Pass an explicit
``font=`` / ``size_pt=`` only where a single run genuinely differs.
"""

import os
import sys

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from style_profiles import active_profile, heading_spec, resolve_profile, set_active_profile
except ImportError:
    # Loaded without the bundle's scripts/ dir on sys.path (importlib spec, exec,
    # a build script that forgot the sys.path.insert). Find our sibling ourselves
    # rather than failing the whole build over an import line.
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from style_profiles import active_profile, heading_spec, resolve_profile, set_active_profile

# Kept for build scripts that name a face directly. These are the "modern"
# profile's values; helpers no longer reference them — they read the active
# profile instead, so a 公文 does not silently come out in 微软雅黑.
CN_FONT = "微软雅黑"
CN_FONT_SERIF = "宋体"
CN_FONT_HEADING = "微软雅黑"
LATIN_FONT = "Times New Roman"

HEADER_FILL = "1F4E79"
BAND_FILL = "F2F6FA"


def set_run_font(run, name: str | None = None, size_pt: float | None = None, bold=None, color: str | None = None):
    """Set a run's font so it applies to Chinese too.

    ``run.font.name`` only writes ``w:ascii`` and ``w:hAnsi``; without a
    matching ``w:eastAsia`` entry Word renders CJK in the theme font and the
    document silently ignores the face you asked for.
    """
    if name:
        run.font.name = name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)
        rfonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _style_font(style, cn_name: str, size_pt: float | None = None):
    style.font.name = cn_name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), cn_name)
    rfonts.set(qn("w:hAnsi"), cn_name)
    rfonts.set(qn("w:eastAsia"), cn_name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)


def _apply_line_spacing(paragraph_format, spec: dict):
    """公文 leading is a fixed point value; the corporate profile uses a multiple.

    A multiple cannot hold GB/T 9704's "22 lines per page" once the body size
    moves, so the two are expressed differently and only one is ever set.
    """
    if spec.get("line_pt"):
        paragraph_format.line_spacing = Pt(spec["line_pt"])
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    elif spec.get("line_multiple"):
        paragraph_format.line_spacing = spec["line_multiple"]


def apply_chinese_defaults(
    doc,
    profile=None,
    body_font: str | None = None,
    heading_font: str | None = None,
    body_pt: float | None = None,
    **overrides,
):
    """Set Normal and Heading 1-4, and make this profile the default for every
    other helper in the package.

    Defaults to the GB/T 9704 公文 profile. For a CV, a marketing one-pager or an
    outward-facing proposal — anything explicitly not an official document — pass
    ``profile="modern"``.

    ``body_font`` / ``heading_font`` / ``body_pt`` stay supported as shorthand for
    a one-off tweak on top of the chosen profile. Note ``heading_font`` rewrites
    **all four levels**: under 公文 that costs the level-2 楷体 and with it the
    "same size, different face" ladder the standard is built on. To change one
    level, pass ``profile={"headings": {1: {"font": ...}}}`` instead.

    Do not call this on a document opened from ``uploads/`` — it rewrites Normal,
    Heading 1–4 and the margins, which erases the styling the user's own file
    came with.

    Word's stock Heading styles are Calibri Light in a blue nobody asked for;
    left alone they make every generated document look identically foreign.
    """
    tweaks: dict = {}
    if body_font:
        tweaks.setdefault("body", {})["font"] = body_font
    if body_pt is not None:
        tweaks.setdefault("body", {})["pt"] = body_pt
    if heading_font:
        tweaks["headings"] = {level: {"font": heading_font} for level in (1, 2, 3, 4)}
    for key, value in overrides.items():
        tweaks[key] = value

    resolved = set_active_profile(profile, tweaks or None)
    body = resolved["body"]

    # Margins belong to the profile too, and this function is usually called
    # *after* setup_page (that is the order SKILL.md's skeleton shows), so a
    # profile switch has to reach back and correct the page it already laid out
    # — otherwise a "modern" document keeps 公文's 37/35/28/26mm 版心.
    # To override margins deliberately, call setup_page(margin_cm=...) *after*
    # this function.
    page = resolved["page"]
    for section in doc.sections:
        section.top_margin = Cm(page["margin_top_cm"])
        section.bottom_margin = Cm(page["margin_bottom_cm"])
        section.left_margin = Cm(page["margin_left_cm"])
        section.right_margin = Cm(page["margin_right_cm"])

    normal = doc.styles["Normal"]
    _style_font(normal, body["font"], body["pt"])
    _apply_line_spacing(normal.paragraph_format, body)
    normal.paragraph_format.space_after = Pt(body.get("space_after_pt", 0))
    if body.get("color"):
        normal.font.color.rgb = RGBColor.from_string(body["color"])

    headings = resolved["headings"]
    for level in (1, 2, 3, 4):
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        spec = heading_spec(level, resolved)
        _style_font(style, spec["font"], spec["pt"])
        style.font.bold = spec.get("bold", True)
        # Word's stock Heading 4 (and several template variants) are italic.
        # Left alone, a level-4 Chinese heading renders as slanted serif and
        # looks like a different document.
        style.font.italic = False
        if spec.get("color"):
            style.font.color.rgb = RGBColor.from_string(spec["color"])
        before = headings.get("space_before_pt", 12)
        style.paragraph_format.space_before = Pt(before + 6 if level == 1 and before else before)
        style.paragraph_format.space_after = Pt(headings.get("space_after_pt", 6))
    return resolved


def apply_gongwen_defaults(doc, **overrides):
    """GB/T 9704 公文 typography. Same as ``apply_chinese_defaults(doc)`` — named
    explicitly for build scripts that want the intent on the page."""
    return apply_chinese_defaults(doc, profile="gongwen", **overrides)


def setup_page(
    doc,
    width_cm: float = 21.0,
    height_cm: float = 29.7,
    margin_cm: float | None = None,
    landscape: bool = False,
    profile=None,
):
    """A4 portrait by default; pass landscape=True to swap the dimensions.

    Margins come from the active profile (公文 uses the standard's uneven
    37/35/28/26mm 版心). Pass ``margin_cm`` to force one value on all four sides.
    """
    page = (resolve_profile(profile) if profile is not None else active_profile())["page"]
    if margin_cm is not None:
        top = bottom = left = right = margin_cm
    else:
        top = page["margin_top_cm"]
        bottom = page["margin_bottom_cm"]
        left = page["margin_left_cm"]
        right = page["margin_right_cm"]
    for section in doc.sections:
        section.page_width = Cm(height_cm if landscape else width_cm)
        section.page_height = Cm(width_cm if landscape else height_cm)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
    return doc.sections[0]


def content_width_cm(section) -> float:
    """Usable width between the margins, in cm — the cap for tables and images."""
    return (section.page_width - section.left_margin - section.right_margin) / 360000


def _field(paragraph, instruction: str, placeholder: str = ""):
    """Insert a Word field code. Fields are how TOC and page numbers work."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._element.append(begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._element.append(instr)

    run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._element.append(sep)

    if placeholder:
        paragraph.add_run(placeholder)

    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._element.append(end)
    return paragraph


def enable_update_fields(doc):
    """Ask Word to refresh every field when the document is opened.

    Without this the TOC shows the placeholder text until someone presses F9,
    which reads as a broken document.
    """
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def add_toc(doc, levels: str = "1-3", title: str = "目录", font: str | None = None, size_pt: float | None = None):
    """Insert a table of contents field. Requires built-in Heading styles."""
    spec = active_profile()["toc"]
    if title:
        heading = doc.add_paragraph()
        run = heading.add_run(title)
        set_run_font(run, font or spec["font"], size_pt or spec["pt"], bold=True, color=spec.get("color"))
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    _field(paragraph, f' TOC \\o "{levels}" \\h \\z \\u ', "（打开文档时自动生成；若未显示请按 Ctrl+A 后 F9）")
    enable_update_fields(doc)
    return paragraph


def add_page_number_footer(
    section,
    template: str = "第 {PAGE} 页 / 共 {NUMPAGES} 页",
    size_pt: float | None = None,
    font: str | None = None,
):
    """Page numbers in the footer, as live fields rather than baked-in text."""
    spec = active_profile()["footer"]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for chunk in _split_template(template):
        if chunk in ("PAGE", "NUMPAGES"):
            _field(paragraph, f" {chunk} ", "1")
        else:
            paragraph.add_run(chunk)
    for run in paragraph.runs:
        set_run_font(run, font or spec["font"], size_pt or spec["pt"], color=spec.get("color"))
    return paragraph


def _split_template(template: str):
    parts, buffer, index = [], "", 0
    while index < len(template):
        if template[index] == "{":
            close = template.find("}", index)
            if close > 0:
                if buffer:
                    parts.append(buffer)
                    buffer = ""
                parts.append(template[index + 1 : close])
                index = close + 1
                continue
        buffer += template[index]
        index += 1
    if buffer:
        parts.append(buffer)
    return parts


def set_cell_shading(cell, hex_fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def add_table(
    doc,
    headers,
    rows,
    widths_cm=None,
    style: str = "Table Grid",
    font_pt: float | None = None,
    banded: bool | None = None,
    font: str | None = None,
    size_pt: float | None = None,
):
    """Add a table with widths Word will actually honour.

    python-docx exposes ``column.width``, but Word ignores it unless the same
    width is written to every cell in that column and autofit is off — the
    single most common "why is my table squashed" bug.

    Face, size and header treatment come from the active profile: the corporate
    profile gives a dark header band with white text, 公文 gives a plain bold
    header row (an official document does not carry colour blocks).
    """
    spec = active_profile()["table"]
    face = font or spec["font"]
    pt = size_pt if size_pt is not None else (font_pt if font_pt is not None else spec["pt"])
    striped = spec.get("banded", True) if banded is None else banded

    table = doc.add_table(rows=1, cols=len(headers))
    if style:
        try:
            table.style = style
        except KeyError:
            pass  # style missing from the template; grid lines are cosmetic
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths_cm is None

    header_cells = table.rows[0].cells
    for index, title in enumerate(headers):
        cell = header_cells[index]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(title))
        set_run_font(run, face, pt, bold=spec.get("header_bold", True), color=spec.get("header_color"))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if spec.get("header_fill"):
            set_cell_shading(cell, spec["header_fill"])

    for row_index, record in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(record):
            cell = cells[index]
            cell.text = ""
            run = cell.paragraphs[0].add_run("" if value is None else str(value))
            set_run_font(run, face, pt)
            if isinstance(value, (int, float)):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if striped and spec.get("band_fill") and row_index % 2 == 1:
                set_cell_shading(cell, spec["band_fill"])

    if widths_cm:
        for row in table.rows:
            for index, width in enumerate(widths_cm):
                if index < len(row.cells):
                    row.cells[index].width = Cm(width)
    repeat_header_row(table)
    return table


def repeat_header_row(table):
    """Repeat row 0 on every page the table spills onto.

    Word only does this when the row carries ``w:tblHeader``; without it a long
    table's second page is a wall of numbers with no column names.
    """
    if not table.rows:
        return table
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))
    return table


def add_hr(doc, color: str = "BFBFBF", size: int = 6):
    """A horizontal rule as a paragraph bottom border.

    Never use a one-row table for this — it breaks text flow and screen readers.
    """
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)
    return paragraph


def add_image_fitted(doc, image_path: str, section=None, max_width_cm: float | None = None, caption: str | None = None):
    """Insert an image capped at the text width, optionally captioned.

    Only ever shrinks: forcing every picture to the full text width blows a
    600px chart up to 16cm and it renders visibly soft.
    """
    if max_width_cm is None:
        max_width_cm = content_width_cm(section or doc.sections[0])
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = paragraph.add_run().add_picture(image_path)
    limit = Cm(max_width_cm)
    if picture.width > limit:
        picture.height = int(picture.height * limit / picture.width)
        picture.width = limit
    if caption:
        spec = active_profile()["caption"]
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cap.add_run(caption), spec["font"], spec["pt"], color=spec.get("color"))
    return paragraph


def first_line_indent(paragraph, chars: float | None = None, font_pt: float | None = None):
    """Chinese body text is indented by two characters, not by a tab."""
    body = active_profile()["body"]
    chars = body.get("indent_chars", 2) if chars is None else chars
    font_pt = body["pt"] if font_pt is None else font_pt
    paragraph.paragraph_format.first_line_indent = Pt(chars * font_pt)
    return paragraph


def add_body(
    doc,
    text: str,
    indent: bool = True,
    font_pt: float | None = None,
    font: str | None = None,
    size_pt: float | None = None,
):
    """A body paragraph with the Chinese conventions already applied."""
    body = active_profile()["body"]
    face = font or body["font"]
    pt = size_pt if size_pt is not None else (font_pt if font_pt is not None else body["pt"])
    paragraph = doc.add_paragraph()
    set_run_font(paragraph.add_run(text), face, pt, color=body.get("color"))
    _apply_line_spacing(paragraph.paragraph_format, body)
    if indent:
        first_line_indent(paragraph, body.get("indent_chars", 2), pt)
    return paragraph


def add_heading_cn(doc, text: str, level: int = 1, font: str | None = None, size_pt: float | None = None):
    """A heading that keeps the built-in style (so the TOC finds it) but renders
    in the active profile's face for that level.

    Under 公文 all four levels are 三号 and differ by face (黑体 / 楷体 / 仿宋加粗)
    rather than by size — do not "fix" that into a descending size ladder.
    """
    spec = heading_spec(level)
    heading = doc.add_heading(level=level)
    run = set_run_font(
        heading.add_run(text),
        font or spec["font"],
        size_pt if size_pt is not None else spec["pt"],
        bold=spec.get("bold", True),
        color=spec.get("color"),
    )
    run.font.italic = False  # stock Heading 4 is italic; override at run level too
    return heading


def add_gongwen_title(doc, text: str, font: str | None = None, size_pt: float | None = None):
    """The document title of an official document (标题) — 二号方正小标宋简体, centred.

    Not a Heading style: GB/T 9704's 标题 sits above the 主送机关 and is not part
    of the heading ladder, so it must not land in the table of contents.
    Long titles wrap on meaning; do not split a word or a book-title mark.
    """
    spec = active_profile()["title"]
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    set_run_font(
        paragraph.add_run(text),
        font or spec["font"],
        size_pt if size_pt is not None else spec["pt"],
        bold=False,  # 小标宋 is already heavy; bolding it renders as a smear
        color=spec.get("color"),
    )
    return paragraph


def add_signature_block(doc, org: str | None = None, date: str | None = None, blank_lines: int = 1):
    """发文机关署名 + 成文日期, set back from the right margin per GB/T 9704.

    The two indents differ and the standard is explicit about it: on a document
    that carries no seal, 发文机关署名 is set back **two** characters and 成文日期
    **four**. Equalising them is the giveaway that a template was eyeballed.

    Leave ``org`` empty when the user has not supplied an authorised issuer —
    never invent a real 机关名称.
    """
    body = active_profile()["body"]
    for _ in range(max(0, blank_lines)):
        doc.add_paragraph()
    paragraphs = []
    for text, back_chars in ((org, 2), (date, 4)):
        if not text:
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(back_chars * body["pt"])
        _apply_line_spacing(paragraph.paragraph_format, body)
        set_run_font(paragraph.add_run(text), body["font"], body["pt"], color=body.get("color"))
        paragraphs.append(paragraph)
    return paragraphs
