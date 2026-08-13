import copy
import re
from io import BytesIO
from typing import IO, Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# Separator between the human-readable node name and the lookup key inside a
# placeholder: ``{{display name|node_id.field}}``. The name is a snapshot taken
# when the variable was inserted -- it exists purely so the template is readable
# and is never used to resolve values (release-contract INV-8).
PLACEHOLDER_DISPLAY_SEPARATOR = "|"


def normalize_placeholder_key(raw: str) -> str:
    """Reduce a placeholder body to the key used for variable lookup.

    ``报告生成|node_a.output`` -> ``node_a.output``; a legacy ``node_a.output``
    is returned as-is. Splitting on the LAST separator keeps node names that
    themselves contain a ``|`` from shifting the split point.

    Callers must keep the ORIGINAL placeholder text as the replacement map key
    -- replacement matches the document verbatim, so normalizing too early
    leaves raw ``{{...}}`` behind in the rendered report.
    """
    if PLACEHOLDER_DISPLAY_SEPARATOR in raw:
        raw = raw.rsplit(PLACEHOLDER_DISPLAY_SEPARATOR, 1)[-1]
    return raw.strip()


class DocxReplacer:
    """
    Docx File Content Replacement Tool
    Supports complex format placeholder replacement, including text, tables, images, titles, etc.
    """

    def __init__(self, template_path: str | IO[bytes]):
        self.template_path = template_path
        self.doc = Document(template_path)
        self.placeholder_pattern = re.compile(r"\{\{([^}]+)\}\}")
        self._init_style()

    def check_style(self, style_name: str, **kwargs):
        all_style = self.doc.styles
        if style_name not in all_style:
            style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            for key, value in kwargs.items():
                setattr(font, key, value)
        else:
            style = all_style[style_name]
        return style

    def _init_style(self):
        self.check_style("Heading 1", size=304800)
        self.check_style("Heading 2", size=254000)
        self.check_style("Heading 3", size=203200)
        self.check_style("Heading 4", size=177800)
        self.check_style("Heading 5", size=152400)
        self.check_style("Heading 6", bold=True)

    def replace_and_save(self, variables: dict[str, list[dict[str, Any]]], output_path: str):
        """
        Replace the placeholders and save the document.

        Args:
        Variables: A dictionary of variables, with the format of {"var_name": [{"type": "text", "content": "..."}]}
        Output_path: The path of the output file
        """
        self._process_paragraphs(self.doc.paragraphs, variables)

        for table in self.doc.tables:
            self._process_table(table, variables)

        for section in self.doc.sections:
            if section.header:
                self._process_paragraphs(section.header.paragraphs, variables)
            if section.footer:
                self._process_paragraphs(section.footer.paragraphs, variables)

        self.doc.save(output_path)

    def _process_table(self, table, variables: dict[str, list[dict[str, Any]]]):
        for row in table.rows:
            for cell in row.cells:
                self._process_paragraphs(cell.paragraphs, variables)
                for nested_table in cell.tables:
                    self._process_table(nested_table, variables)

    def _process_paragraphs(self, paragraphs: list[Paragraph], variables: dict[str, list[dict[str, Any]]]):
        i = 0
        while i < len(paragraphs):
            paragraph = paragraphs[i]
            text = paragraph.text
            matches = list(self.placeholder_pattern.finditer(text))

            if matches:
                # Plain-text values are written into the run that holds the
                # placeholder, so the author's underline / font / size survive and
                # the paragraph itself (style, numbering, borders) is never
                # rebuilt. Anything block-level still takes the rebuild path.
                if not (
                    self._placeholders_are_inline(matches, variables)
                    and self._replace_inline_placeholders(paragraph, variables)
                ):
                    insert_index = self._get_paragraph_index(paragraph)
                    self._replace_paragraph_placeholders(paragraph, matches, variables, insert_index)

            i += 1

    def _placeholders_are_inline(
        self,
        matches: list[re.Match],
        variables: dict[str, list[dict[str, Any]]],
    ) -> bool:
        """Whether every placeholder in this paragraph resolves to plain text.

        A table, image or heading cannot live inside a run — those paragraphs
        have to be split apart, which is what the rebuild path does.
        """
        for match in matches:
            items = variables.get(match.group(1))
            if items is None:
                # Unresolved placeholder: left verbatim either way.
                continue
            if any(item.get("type") != "text" for item in items):
                return False
        return True

    def _replace_inline_placeholders(
        self,
        paragraph: Paragraph,
        variables: dict[str, list[dict[str, Any]]],
    ) -> bool:
        """Rewrite text placeholders in place, one run at a time.

        Returns False when the placeholders cannot be located in the run text —
        a placeholder split across a hyperlink, say — so the caller can fall
        back to rebuilding the paragraph.
        """
        runs = paragraph.runs
        if not runs:
            return False

        runs_text = "".join(run.text for run in runs)
        matches = list(self.placeholder_pattern.finditer(runs_text))
        if not matches:
            return False

        # Right to left: rewriting a later span never shifts an earlier offset.
        for match in reversed(matches):
            items = variables.get(match.group(1))
            if items is None:
                continue
            self._replace_run_span(paragraph, match.start(), match.end(), items)
        return True

    def _replace_run_span(
        self,
        paragraph: Paragraph,
        start: int,
        end: int,
        items: list[dict[str, Any]],
    ):
        """Replace the characters [start, end) of a paragraph's run text."""
        spans = []
        offset = 0
        for run in paragraph.runs:
            spans.append((offset, offset + len(run.text), run))
            offset += len(run.text)

        touched = [
            (run_start, run_end, run) for run_start, run_end, run in spans if run_end > start and run_start < end
        ]
        if not touched:
            return

        first_start, _, first_run = touched[0]
        last_start, _, last_run = touched[-1]
        prefix = first_run.text[: start - first_start]
        suffix = last_run.text[end - last_start :]

        # Snapshot the placeholder run's formatting before writing to it: extra
        # value items and the trailing text should inherit what the author put on
        # the placeholder, not what the first item asks for.
        template_element = copy.deepcopy(first_run._element)

        for _run_start, _run_end, run in touched[1:]:
            run.text = ""

        text_items = items or [{"content": ""}]
        first_run.text = prefix + text_items[0].get("content", "")
        self._apply_run_format(first_run, text_items[0])

        anchor = first_run._element
        for item in text_items[1:]:
            anchor = self._insert_run_after(paragraph, anchor, template_element, item.get("content", ""), item)

        if suffix:
            if last_run is not first_run:
                last_run.text = suffix
            else:
                self._insert_run_after(paragraph, anchor, template_element, suffix, {})

    def _insert_run_after(
        self,
        paragraph: Paragraph,
        anchor_element,
        template_element,
        text: str,
        format_data: dict[str, Any],
    ):
        new_element = copy.deepcopy(template_element)
        anchor_element.addnext(new_element)
        run = Run(new_element, paragraph)
        run.text = text
        self._apply_run_format(run, format_data)
        return new_element

    def _get_paragraph_index(self, paragraph: Paragraph) -> int:
        parent = paragraph._element.getparent()
        return parent.index(paragraph._element)

    def _replace_paragraph_placeholders(
        self,
        paragraph: Paragraph,
        matches: list[re.Match],
        variables: dict[str, list[dict[str, Any]]],
        insert_index: int,
    ):
        parent = paragraph._element.getparent()
        text = paragraph.text

        segments = []
        last_end = 0

        for match in matches:
            var_name = match.group(1)
            start, end = match.span()

            if start > last_end:
                segments.append({"type": "text_segment", "content": text[last_end:start], "paragraph": paragraph})

            if var_name in variables:
                segments.append({"type": "variable", "content": variables[var_name], "paragraph": paragraph})
            else:
                segments.append({"type": "text_segment", "content": match.group(0), "paragraph": paragraph})

            last_end = end

        if last_end < len(text):
            segments.append({"type": "text_segment", "content": text[last_end:], "paragraph": paragraph})

        original_format = self._extract_paragraph_format(paragraph)
        original_run_format = self._extract_run_format(paragraph.runs[0] if paragraph.runs else None)

        parent.remove(paragraph._element)

        current_insert_index = insert_index
        current_paragraph = None

        for segment in segments:
            if segment["type"] == "text_segment":
                if current_paragraph is None:
                    current_paragraph = self._insert_paragraph_at_index(parent, current_insert_index, original_format)
                    current_insert_index += 1

                run = current_paragraph.add_run(segment["content"])
                self._apply_run_format(run, original_run_format)

            elif segment["type"] == "variable":
                for item in segment["content"]:
                    item_type = item.get("type")

                    if item_type == "text":
                        if current_paragraph is None:
                            current_paragraph = self._insert_paragraph_at_index(
                                parent, current_insert_index, original_format
                            )
                            current_insert_index += 1

                        run = current_paragraph.add_run(item["content"])
                        self._apply_run_format(run, item)

                    elif item_type in ["table", "image", "heading"]:
                        if current_paragraph is not None and current_paragraph.text.strip():
                            current_paragraph = None

                        if item_type == "table":
                            self._insert_table_at_index(parent, current_insert_index, item)
                        elif item_type == "image":
                            self._insert_image_at_index(parent, current_insert_index, item, original_format)
                        elif item_type == "heading":
                            self._insert_heading_at_index(parent, current_insert_index, item)

                        current_insert_index += 1
                        current_paragraph = None

    def _extract_paragraph_format(self, paragraph: Paragraph) -> dict[str, Any]:
        return {
            "alignment": paragraph.alignment,
            "left_indent": paragraph.paragraph_format.left_indent,
            "right_indent": paragraph.paragraph_format.right_indent,
            "first_line_indent": paragraph.paragraph_format.first_line_indent,
            "space_before": paragraph.paragraph_format.space_before,
            "space_after": paragraph.paragraph_format.space_after,
            "line_spacing": paragraph.paragraph_format.line_spacing,
        }

    def _extract_run_format(self, run) -> dict[str, Any]:
        if run is None:
            return {}

        return {
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_name": run.font.name,
            "font_size": run.font.size,
            "font_color": run.font.color.rgb if run.font.color.rgb else None,
        }

    def _insert_paragraph_at_index(self, parent, index: int, format_dict: dict[str, Any]) -> Paragraph:
        p_element = OxmlElement("w:p")
        parent.insert(index, p_element)
        paragraph = Paragraph(p_element, self.doc)

        # 应用格式
        if format_dict.get("alignment") is not None:
            paragraph.alignment = format_dict["alignment"]
        if format_dict.get("left_indent") is not None:
            paragraph.paragraph_format.left_indent = format_dict["left_indent"]
        if format_dict.get("right_indent") is not None:
            paragraph.paragraph_format.right_indent = format_dict["right_indent"]
        if format_dict.get("first_line_indent") is not None:
            paragraph.paragraph_format.first_line_indent = format_dict["first_line_indent"]
        if format_dict.get("space_before") is not None:
            paragraph.paragraph_format.space_before = format_dict["space_before"]
        if format_dict.get("space_after") is not None:
            paragraph.paragraph_format.space_after = format_dict["space_after"]
        if format_dict.get("line_spacing") is not None:
            paragraph.paragraph_format.line_spacing = format_dict["line_spacing"]

        return paragraph

    def _apply_run_format(self, run, format_data: dict[str, Any]):
        if format_data.get("bold"):
            run.bold = True
        if format_data.get("italic"):
            run.italic = True
        if format_data.get("underline"):
            run.underline = True
        if format_data.get("font_size"):
            if isinstance(format_data["font_size"], int):
                run.font.size = Pt(format_data["font_size"])
            else:
                run.font.size = format_data["font_size"]
        if format_data.get("font_name"):
            run.font.name = format_data["font_name"]
        if format_data.get("color"):
            if isinstance(format_data["color"], tuple) and len(format_data["color"]) == 3:
                run.font.color.rgb = RGBColor(*format_data["color"])
        if format_data.get("font_color"):
            run.font.color.rgb = format_data["font_color"]

    def _insert_table_at_index(self, parent, index: int, item: dict[str, Any]):
        data = item["content"]
        rows = len(data)
        cols = len(data[0]) if rows > 0 else 0

        if rows == 0 or cols == 0:
            return

        table = self.doc.add_table(rows=0, cols=cols)

        if item.get("style"):
            table.style = item["style"]

        for row_data in data:
            row = table.add_row()
            for col_idx, cell_content in enumerate(row_data):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    self._fill_cell(cell, cell_content)

        tbl_element = table._element
        parent.remove(tbl_element)
        parent.insert(index, tbl_element)

    def _fill_cell(self, cell: _Cell, cell_content: Any):
        """
        Filling Cell Content

        Cell content format:
            1. Single element (dictionary): {"type": "text", "content": "..."}
            2. Composite elements (list): [{"type": "text", ...}, {"type": "image", ...}]

        Args:
            cell: The cell object
            cell_content: The cell content, a dictionary or a list
        """
        if cell.paragraphs:
            default_paragraph = cell.paragraphs[0]
            for run in default_paragraph.runs:
                run.text = ""
        else:
            default_paragraph = cell.add_paragraph()

        if isinstance(cell_content, dict):
            if "type" not in cell_content or "content" not in cell_content:
                raise ValueError(
                    f"The cell element must contain the `type` and `content` fields, but got：{cell_content}"
                )
            cell_content = [cell_content]
        elif isinstance(cell_content, list):
            for element in cell_content:
                if not isinstance(element, dict) or "type" not in element or "content" not in element:
                    raise ValueError(
                        f"The cell element must contain the `type` and `content` fields, but got：{element}"
                    )
        else:
            raise ValueError(f"Not supported data type：{type(cell_content)}")

        current_paragraph = default_paragraph

        for element in cell_content:
            element_type = element["type"]
            if element_type == "text":
                run = current_paragraph.add_run(element["content"])
                self._apply_run_format(run, element)

            elif element_type == "image":
                if current_paragraph.text.strip():
                    current_paragraph = cell.add_paragraph()

                self._add_image_to_paragraph(current_paragraph, element)
                current_paragraph = cell.add_paragraph()

            elif element_type == "paragraph":
                current_paragraph = cell.add_paragraph()
                if element.get("alignment"):
                    current_paragraph.alignment = element["alignment"]

                if isinstance(element["content"], str):
                    run = current_paragraph.add_run(element["content"])
                    self._apply_run_format(run, element)
                elif isinstance(element["content"], list):
                    for text_item in element["content"]:
                        if not isinstance(text_item, dict) or "type" not in text_item:
                            raise ValueError(
                                f"Paragraph content elements must include a `type` field; got：{text_item}"
                            )
                        if text_item["type"] == "text":
                            run = current_paragraph.add_run(text_item["content"])
                            self._apply_run_format(run, text_item)
            if element.get("alignment"):
                for cell_paragraph in cell.paragraphs:
                    cell_paragraph.alignment = element["alignment"]

    def _add_image_to_paragraph(self, paragraph: Paragraph, image_data: dict[str, Any]):
        run = paragraph.add_run()

        try:
            width = Inches(image_data.get("width", 2))
            height = Inches(image_data.get("height")) if image_data.get("height") else None

            if isinstance(image_data["content"], str):
                # local file path
                if height:
                    run.add_picture(image_data["content"], width=width, height=height)
                else:
                    run.add_picture(image_data["content"], width=width)
            elif isinstance(image_data["content"], bytes):
                # bytes data
                image_stream = BytesIO(image_data["content"])
                if height:
                    run.add_picture(image_stream, width=width, height=height)
                else:
                    run.add_picture(image_stream, width=width)
        except Exception as e:
            paragraph.add_run(f"Image add failed: {e!s}]")

        # set alignment
        if image_data.get("alignment"):
            paragraph.alignment = image_data["alignment"]

    def _insert_image_at_index(self, parent, index: int, item: dict[str, Any], paragraph_format: dict[str, Any]):
        paragraph = self._insert_paragraph_at_index(parent, index, paragraph_format)
        self._add_image_to_paragraph(paragraph, item)

    def _insert_heading_at_index(self, parent, index: int, item: dict[str, Any]):
        p_element = OxmlElement("w:p")
        parent.insert(index, p_element)
        paragraph = Paragraph(p_element, self.doc)

        level = item.get("level", 1)
        paragraph.style = f"Heading {level}"

        run = paragraph.add_run(item["content"])
        self._apply_run_format(run, item)

    def extract_variables(self) -> list[str]:
        variables = []
        seen = set()

        for paragraph in self.doc.paragraphs:
            vars_in_paragraph = self._extract_vars_from_text(paragraph.text)
            for var in vars_in_paragraph:
                if var not in seen:
                    variables.append(var)
                    seen.add(var)

        for table in self.doc.tables:
            vars_in_table = self._extract_vars_from_table(table)
            for var in vars_in_table:
                if var not in seen:
                    variables.append(var)
                    seen.add(var)

        for section in self.doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    vars_in_paragraph = self._extract_vars_from_text(paragraph.text)
                    for var in vars_in_paragraph:
                        if var not in seen:
                            variables.append(var)
                            seen.add(var)

            if section.footer:
                for paragraph in section.footer.paragraphs:
                    vars_in_paragraph = self._extract_vars_from_text(paragraph.text)
                    for var in vars_in_paragraph:
                        if var not in seen:
                            variables.append(var)
                            seen.add(var)

        return variables

    def _extract_vars_from_text(self, text: str) -> list[str]:
        matches = self.placeholder_pattern.findall(text)
        return matches

    def _extract_vars_from_table(self, table) -> list[str]:
        variables = []
        seen = set()

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    vars_in_paragraph = self._extract_vars_from_text(paragraph.text)
                    for var in vars_in_paragraph:
                        if var not in seen:
                            variables.append(var)
                            seen.add(var)
                for nested_table in cell.tables:
                    nested_vars = self._extract_vars_from_table(nested_table)
                    for var in nested_vars:
                        if var not in seen:
                            variables.append(var)
                            seen.add(var)

        return variables
