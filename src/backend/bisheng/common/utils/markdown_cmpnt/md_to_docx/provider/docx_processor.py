# noinspection PyProtectedMember
#
import io
import re
from socket import socket
from urllib.error import HTTPError, URLError
from urllib.request import urlopen

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import *
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, RGBColor, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from bisheng.common.utils.markdown_cmpnt.md_to_docx.provider.docx_plus import add_hyperlink
from bisheng.common.utils.markdown_cmpnt.md_to_docx.provider.style_manager import StyleManager
from bisheng.common.utils.markdown_cmpnt.md_to_docx.utils.style_enum import MDX_STYLE


class DocxProcessor:
    def __init__(self, style_conf: dict, debug_state: bool = False, show_image_desc: bool = True):
        """
        InisialisasiDocxProcessor
        :param style_conf: Style Configuration Dictionary
        :param debug_state: Whether to turn on debug mode
        :param show_image_desc: Whether to display the description of the picture, i.e. `![desc](src/img)` II descThe contents of the
        """
        self.document = Document()
        self.debug_state = debug_state
        self.show_image_desc = show_image_desc
        if style_conf is not None:
            StyleManager(self.document, style_conf).init_styles()

    def debug(self, *args):
        """Debug Output"""
        if self.debug_state:
            print(*args)

    # h1, h2, ...
    def add_heading(self, content: str, tag: str):
        level: int = int(tag.__getitem__(1))
        p = self.document.add_paragraph(content, style="Heading%d" % level)
        # Force title not pagination
        p.paragraph_format.page_break_before = False
        p.paragraph_format.keep_with_next = True
        return p

    def add_run(self, p: Paragraph, content: str, char_style: str = "plain"):
        # fixme Sentences with more than one style in a row are ignored, such as:
        # <u>**Bold and*Italic*Underline again**</u>
        self.debug("[%s]:" % char_style, content)
        run = p.add_run(content)

        # Standardized label name - will be HTML5 Label mapping to standard styles
        style_map = {
            'b': 'strong',      # <b> -> Bold
            'i': 'em',          # <i> -> Italic
            'del': 'strike',    # <del> -> Strikethrough
            'mark': 'highlight' # <mark> -> Gao Liang
        }
        char_style = style_map.get(char_style, char_style)

        # Should not be used in the form of run.bold = (char_style=="strong") to be
        # Because there is no explicit bolding, it does not mean that the whole paragraph is not bold.
        if char_style == "strong":
            run.bold = True
        if char_style == "em":
            run.italic = True
        if char_style == "u":
            run.underline = True
        if char_style == "strike":
            run.font.strike = True
        if char_style == "sub":
            run.font.subscript = True
        if char_style == "sup":
            run.font.superscript = True
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW if char_style == "highlight" else None

        # if char_style == "code":
        #     run.font.name = "Consolas"

    def add_code_block(self, pre_tag):
        # TODO Code-Block Styles
        # TODO Set the Chinese font in the code block (table), it seems that it can only be specified by Chinese fonts are styled to get there.
        code_table = self.document.add_table(0, 1, style=MDX_STYLE.TABLE)
        row_cells = code_table.add_row().cells

        # Security check: Make sure the code block has content
        if pre_tag.contents and len(pre_tag.contents) > 0 and pre_tag.contents[0].string:
            code_text = pre_tag.contents[0].string.rstrip('\n')
            run = row_cells[0].paragraphs[0].add_run(code_text)
            run.font.name = "Consolas"
        else:
            # If the code block is empty, add a blank placeholder
            run = row_cells[0].paragraphs[0].add_run("")
            run.font.name = "Consolas"

    def add_picture(self, img_tag, parent_paragraph: Paragraph = None):
        """
        Adding Images to Documents
        :param img_tag: Images Tab
        :param parent_paragraph: Parent paragraph. If provided,Images will be embedded in the paragraph;Otherwise create a new standalone paragraph
        """
        # If no parent paragraph is provided,Create a new standalone paragraph(Appear from center)
        if parent_paragraph is None:
            p: Paragraph = self.document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.first_line_indent = 0
        else:
            # Use provided parent paragraph(Embedded Mode)
            p = parent_paragraph

        run: Run = p.add_run()

        img_src: str
        scale: float = 100  # Highest priority, unit %
        width_px: int = 100
        height_px: int = 100

        # Set width
        if img_tag.get("style"):
            style_content: str = img_tag["style"]
            img_attr: list = style_content.strip().split(";")
            # print(img_attr)
            attr: str
            for attr in img_attr:
                if attr.find("width") != -1:
                    # TODO <g id="Bold">Medical Treatment:</g> style Medium Width and Height Properties
                    width_px = int(re.findall(r"\d+", attr)[0])
                if attr.find("height") != -1:
                    height_px = int(re.findall(r"\d+", attr)[0])
                if attr.find("zoom") != -1:
                    scale = int(re.findall(r"\d+", attr)[0])

        if img_tag["src"] != "":
            img_src = img_tag["src"]
            # webauthn
            if img_src.startswith("http://") or img_src.startswith("https://"):
                print("[IMAGE] fetching:", img_src)
                try:
                    image_bytes = urlopen(img_src, timeout=10).read()
                    data_stream = io.BytesIO(image_bytes)
                    run.add_picture(data_stream, width=Inches(5.7 * scale / 100))
                except HTTPError as e:
                    print(f"[HTTP ERROR] {e.code}: {img_src}")
                except socket.timeout:
                    print(f"[TIMEOUT] Image load timeout: {img_src}")
                except URLError as e:
                    print(f"[URL ERROR] Failed to fetch image: {e.reason} - {img_src}")
                except Exception as e:
                    print(f"[RESOURCE ERROR] {type(e).__name__}: {e} - {img_src}")
            else:
                # Location Image
                try:
                    run.add_picture(img_src, width=Inches(5.7 * scale / 100))
                except FileNotFoundError:
                    print(f"[FILE ERROR] Image not found: {img_src}")
                except Exception as e:
                    print(f"[RESOURCE ERROR] Failed to load image: {type(e).__name__}: {e} - {img_src}")
        else:
            # webauthn
            img_src = img_tag["title"]
            print("[IMAGE] fetching:", img_src)
            try:
                image_bytes = urlopen(img_src, timeout=10).read()
                data_stream = io.BytesIO(image_bytes)
                run.add_picture(data_stream, width=Inches(5.7 * scale / 100))
            except HTTPError as e:
                print(f"[HTTP ERROR] {e.code}: {img_src}")
            except socket.timeout:
                print(f"[TIMEOUT] Image load timeout: {img_src}")
            except URLError as e:
                print(f"[URL ERROR] Failed to fetch image: {e.reason} - {img_src}")
            except Exception as e:
                print(f"[RESOURCE ERROR] {type(e).__name__}: {e} - {img_src}")

        # If you choose to display an image description, the description will appear below the image
        # Note: Add descriptions only in standalone paragraph mode(Avoid interrupting paragraph flow)
        if parent_paragraph is None and self.show_image_desc and img_tag.get("alt"):
            # TODO Display style for image description
            desc: Paragraph = self.document.add_paragraph(img_tag["alt"], style=MDX_STYLE.CAPTION)
            desc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            desc.style.font.color.rgb = RGBColor(11, 11, 11)
            desc.style.font.bold = False
            desc.paragraph_format.first_line_indent = 0

    def add_table(self, table_root):
        # Number of statistical columns - Preferred From thead Fetch, otherwise from tbody Get on the first line
        col_count: int = 0
        has_thead = False

        if table_root.thead and table_root.thead.tr:
            has_thead = True
            for col in table_root.thead.tr.contents:
                if col.name in ['th', 'td']:
                    col_count += 1
        elif table_root.tbody:
            # FROM tbody Number of columns in the first row
            first_row = table_root.tbody.find('tr')
            if first_row:
                for col in first_row.contents:
                    if col.name in ['th', 'td']:
                        col_count += 1

        if col_count == 0:
            # Fault tolerance: if the number of columns is not found, the default is1
            col_count = 1

        table = self.document.add_table(0, col_count, style=MDX_STYLE.TABLE)  # TODO Table Style

        # Table header row (if present)
        if has_thead:
            head_row_cells = table.add_row().cells
            i = 0
            for col in table_root.thead.tr.contents:
                if col.name in ['th', 'td']:
                    cell_text = col.get_text(strip=True)
                    head_row_cells[i].paragraphs[0].add_run(cell_text).bold = True
                    i += 1

        # Data Line
        if table_root.tbody:
            for tr in table_root.tbody:
                if tr.name != 'tr':
                    continue
                row_cells = table.add_row().cells
                i = 0
                for td in tr.contents:
                    if td.name in ['th', 'td']:
                        cell_text = td.get_text(strip=True)
                        row_cells[i].text = cell_text
                        i += 1

    def add_number_list(self, number_list):
        # print(number_list.contents, "\n")
        for item in number_list.children:
            if item.name != 'li':  # Skip non- li label
                continue

            # Get direct text content(Exclude subtags)
            direct_text = ''.join([str(s) for s in item.find_all(text=True, recursive=False)]).strip()

            # Check if there are paragraph tags
            has_paragraph_tag = item.find('p') is not None

            # If there is direct text,Add your paragraph
            if direct_text:
                self.add_paragraph(item, p_style=MDX_STYLE.LIST_NUMBER) \
                    .style.paragraph_format.space_after = Pt(1)  # TODO Number List Style
            # If there is no direct text but there is <p> label,<g id="Bold">Medical Treatment:</g> <p> label
            elif has_paragraph_tag:
                para_count = 0
                for child in item.children:
                    if hasattr(child, 'name') and child.name == 'p':
                        if para_count == 0:
                            # first <p> Use list style
                            self.add_paragraph(child, p_style=MDX_STYLE.LIST_NUMBER) \
                                .style.paragraph_format.space_after = Pt(1)
                        else:
                            # Additional <p> Use continuation style
                            self.add_paragraph(child, p_style=MDX_STYLE.LIST_CONTINUE) \
                                .style.paragraph_format.space_after = Pt(1)
                        para_count += 1

            # Working with Nested Ordered Lists
            if hasattr(item, "ol") and item.ol is not None:
                sub_num: int = 1  # Sub-Serial Number
                for item2 in item.ol.children:
                    if item2.name != 'li':
                        continue
                    self.add_paragraph(item2, prefix="(%d). " % sub_num, p_style=MDX_STYLE.LIST_CONTINUE) \
                        .style.paragraph_format.first_line_indent = 0  # TODO Number List Style
                    sub_num += 1

            # Working with Nested Unordered Lists
            if hasattr(item, "ul") and item.ul is not None:
                for item2 in item.ul.children:
                    if item2.name != 'li':
                        continue
                    self.add_paragraph(item2, prefix="•  ", p_style=MDX_STYLE.LIST_CONTINUE) \
                        .style.paragraph_format.space_after = Pt(1)

    def add_bullet_list(self, bullet_list):
        # It may be. list
        text = str(bullet_list.contents[1].string).strip()
        if text.startswith("[ ]") or text.startswith("[x]"):
            self.add_todo_list(bullet_list)
            return
        for item in bullet_list.children:
            if item.name != 'li':  # Skip non- li label
                continue

            # Get direct text content(Exclude subtags)
            direct_text = ''.join([str(s) for s in item.find_all(text=True, recursive=False)]).strip()

            # Check if there are paragraph tags
            has_paragraph_tag = item.find('p') is not None

            # If there is direct text,Add your paragraph
            if direct_text:
                self.add_paragraph(item, p_style=MDX_STYLE.LIST_BULLET) \
                    .style.paragraph_format.space_after = Pt(1)
            # If there is no direct text but there is <p> label,<g id="Bold">Medical Treatment:</g> <p> label
            elif has_paragraph_tag:
                para_count = 0
                for child in item.children:
                    if hasattr(child, 'name') and child.name == 'p':
                        if para_count == 0:
                            # first <p> Use list style
                            self.add_paragraph(child, p_style=MDX_STYLE.LIST_BULLET) \
                                .style.paragraph_format.space_after = Pt(1)
                        else:
                            # Additional <p> Use continuation style
                            self.add_paragraph(child, p_style=MDX_STYLE.LIST_CONTINUE) \
                                .style.paragraph_format.space_after = Pt(1)
                        para_count += 1

            # <g id="Bold">Medical Treatment:</g> ul Above listings
            if hasattr(item, "ul") and item.ul is not None:
                for item2 in item.ul.children:
                    if item2.string == "\n" or (not item2.string):
                        continue
                    self.add_paragraph(item2, prefix="•  ", p_style=MDX_STYLE.LIST_CONTINUE) \
                        .style.paragraph_format.space_after = Pt(1)

            # <g id="Bold">Medical Treatment:</g> ol Above listings
            if hasattr(item, "ol") and item.ol is not None:
                # read out start Property,If not present, defaults to1
                sub_num = int(item.ol.get('start', 1))
                for item2 in item.ol.children:
                    if item2.name != 'li':
                        continue
                    self.add_paragraph(item2, prefix=f"{sub_num}. ", p_style=MDX_STYLE.LIST_CONTINUE) \
                        .style.paragraph_format.space_after = Pt(1)
                    sub_num += 1

    # falseTODO list
    def add_todo_list(self, todo_list):
        # list_para.style.font.name = "Consolas"
        for item in todo_list.children:
            if item.string == "\n" or (not item.string):
                continue
            text: str = item.string
            list_para = self.document.add_paragraph(style=MDX_STYLE.PLAIN_LIST)
            if text.startswith("[x]"):
                list_para.add_run("[ √ ]").font.name = "Consolas"
                list_para.add_run(text.replace("[x]", " ", 1))
            elif text.startswith("[ ]"):
                list_para.add_run("[   ]").font.name = "Consolas"
                list_para.add_run(text.replace("[ ]", " ", 1))

    # Split Lines
    def add_split_line(self):
        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        border_elm = parse_xml(
            r'<w:pBdr {}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>'.format(nsdecls('w')))
        p._p.get_or_add_pPr().append(border_elm)

    # Hyperlinks
    def add_link(self, p: Paragraph, text: str, href: str):
        self.debug("[link]:", text, "[href]:", href)
        add_hyperlink(p, href, text)
        # run = p.add_run(text)

    def add_paragraph(self, children, p_style: str = None, prefix: str = ""):
        """
        children: list|str
        An element (including an image) within a paragraph. Divided according to whether there is a style, forming a list.
        There are style text such as bold, italics, images, etc.
        As`I am plain _while_ he is **bold**`will be converted to:
        ["I am plain", "while", "he is", "bold"]
        """
        p = self.document.add_paragraph(prefix, style=p_style)
        if type(children) == str:
            p.add_run(children)
            return p
        for elem in children.contents:  # Traverse all elements in a paragraph
            if elem.name == "a":
                self.add_link(p, elem.string, elem["href"])
            elif elem.name == "img":
                self.add_picture(elem, parent_paragraph=p)  # Pass the current paragraph
            elif elem.name is not None:  # Substring with character style
                self.add_run(p, elem.string, elem.name)
            elif not elem.string == "\n":  # Substring without character style
                self.add_run(p, elem)
        return p

    # from docx.enum.style import WD_STYLE
    def add_blockquote(self, children):
        # TODO Place Blockquote on1x1In the table, optimize the display effect of the reference block
        #  Set left indent, up and down spacing
        table: Table = self.document.add_table(0, 1)
        row_cells = table.add_row().cells

        # Supports multiple paragraphs
        para_count = 0
        for child in children.contents:
            # Skip line breaks
            if isinstance(child, str) and child.strip() == "":
                continue

            # Working with paragraph tags
            if hasattr(child, 'name') and child.name == 'p':
                if para_count > 0:
                    # Add a new passage
                    row_cells[0].add_paragraph()
                p = row_cells[0].paragraphs[para_count]

                for elem in child.contents:  # Traverse all elements in a paragraph
                    if elem.name == "a":
                        self.add_link(p, elem.string, elem["href"])
                    elif elem.name == "img":
                        self.add_picture(elem, parent_paragraph=p)  # Pass the current paragraph
                    elif elem.name is not None:  # Substring with character style
                        self.add_run(p, elem.string, elem.name)
                    elif elem.string and elem.string != "\n":  # Substring without character style
                        self.add_run(p, elem.string)
                para_count += 1
            # Working with Direct Text Nodes
            elif isinstance(child, str) and child.strip():
                if para_count == 0:
                    p = row_cells[0].paragraphs[0]
                    para_count = 1
                else:
                    p = row_cells[0].add_paragraph()
                    para_count += 1
                p.add_run(child.strip())

        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="efefef"/>'.format(nsdecls('w')))
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

    def html2docx(self, html_str: str):
        # OpenHTML

        soup = BeautifulSoup(html_str, 'html.parser')

        # Find it securely body label
        body_tag = soup.find('body')
        if body_tag is None:
            # If no evidence of   microbial body,Use entire soup
            body_tag = soup

        # Title
        title_text = ""
        # parse the labels one by one and writewordII
        for root in body_tag.children:
            # Skip plain text nodes and line breaks
            if isinstance(root, str):
                if root.strip():
                    # Working with Direct Text Nodes
                    self.document.add_paragraph(root.strip(), style=MDX_STYLE.PLAIN_TEXT)
                continue

            # debug("<%s>" % root.name)
            if root.name == "p":  # Normal paragraph
                self.add_paragraph(root, p_style=MDX_STYLE.PLAIN_TEXT)
            elif root.name == "blockquote":  # Blockquote
                self.add_blockquote(root)
            elif root.name == "ol":  # Numbered
                self.add_number_list(root)
            elif root.name == "ul":  # Unordered List OR List
                self.add_bullet_list(root)
            elif root.name == "table":  # Table Filter
                self.add_table(root)
            elif root.name == "hr":
                self.add_split_line()
            elif root.name == "pre":
                self.add_code_block(root)
            elif root.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                if title_text == "":
                    title_text = root.string or ""
                self.add_heading(root.string or "", root.name)

        docx_bytes = io.BytesIO()

        self.document.save(docx_bytes)

        docx_bytes.seek(0)

        return docx_bytes.getbuffer(), title_text
