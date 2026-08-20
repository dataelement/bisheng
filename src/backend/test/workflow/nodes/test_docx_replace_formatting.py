"""Report placeholders must keep the formatting the template author gave them.

The replacement used to rebuild the whole paragraph from its plain text, which
dropped every run-level attribute (underline, bold, font) and the paragraph's own
style. Text values are now written into the run that holds the placeholder.
"""

from io import BytesIO

from docx import Document

from bisheng.workflow.nodes.report.docx_replace import DocxReplacer

PLACEHOLDER = "{{输入|input_bfa69.user_input}}"
KEY = "输入|input_bfa69.user_input"


def _render(build_template, variables) -> Document:
    """Run a template through the replacer and return the rendered document."""
    template = BytesIO()
    doc = Document()
    build_template(doc)
    doc.save(template)

    rendered = BytesIO()
    DocxReplacer(BytesIO(template.getvalue())).replace_and_save(variables, rendered)
    rendered.seek(0)
    return Document(rendered)


def test_placeholder_run_formatting_survives():
    def build(doc):
        paragraph = doc.add_paragraph()
        paragraph.add_run("xxxxx. ")
        underlined = paragraph.add_run(PLACEHOLDER)
        underlined.underline = True
        underlined.bold = True
        paragraph.add_run(" xxxxx")

    rendered = _render(build, {KEY: [{"type": "text", "content": "有下划线"}]})
    runs = rendered.paragraphs[0].runs

    assert [run.text for run in runs] == ["xxxxx. ", "有下划线", " xxxxx"]
    assert runs[1].underline is True
    assert runs[1].bold is True
    # The neighbours are untouched — they used to be re-stamped with run[0]'s format.
    assert runs[0].underline is None
    assert runs[2].underline is None


def test_placeholder_split_across_runs_is_replaced():
    """Word routinely splits a typed placeholder over several runs."""

    def build(doc):
        paragraph = doc.add_paragraph()
        for chunk in ("{{输入|input_", "bfa69.user", "_input}}"):
            run = paragraph.add_run(chunk)
            run.underline = True

    rendered = _render(build, {KEY: [{"type": "text", "content": "拼接占位符"}]})
    paragraph = rendered.paragraphs[0]

    assert paragraph.text == "拼接占位符"
    assert next(run for run in paragraph.runs if run.text).underline is True


def test_paragraph_style_survives():
    def build(doc):
        paragraph = doc.add_paragraph(style="Quote")
        paragraph.add_run(PLACEHOLDER)

    rendered = _render(build, {KEY: [{"type": "text", "content": "引用内容"}]})

    assert rendered.paragraphs[0].style.name == "Quote"
    assert rendered.paragraphs[0].text == "引用内容"


def test_value_items_inherit_placeholder_format_and_add_their_own():
    def build(doc):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(PLACEHOLDER)
        run.underline = True

    rendered = _render(
        build,
        {
            KEY: [
                {"type": "text", "content": "普通"},
                {"type": "text", "content": "加粗", "bold": True},
            ]
        },
    )
    runs = [run for run in rendered.paragraphs[0].runs if run.text]

    assert [run.text for run in runs] == ["普通", "加粗"]
    assert all(run.underline is True for run in runs)
    assert runs[1].bold is True
    assert runs[0].bold is None


def test_surrounding_text_in_the_same_run_is_kept():
    def build(doc):
        run = doc.add_paragraph().add_run(f"前缀{PLACEHOLDER}后缀")
        run.underline = True

    rendered = _render(build, {KEY: [{"type": "text", "content": "中间"}]})
    paragraph = rendered.paragraphs[0]

    assert paragraph.text == "前缀中间后缀"
    assert all(run.underline is True for run in paragraph.runs if run.text)


def test_unresolved_placeholder_is_left_verbatim():
    def build(doc):
        doc.add_paragraph().add_run(f"保留 {PLACEHOLDER}")

    rendered = _render(build, {"other": [{"type": "text", "content": "x"}]})

    assert rendered.paragraphs[0].text == f"保留 {PLACEHOLDER}"


def test_block_content_still_splits_the_paragraph():
    """Tables cannot live inside a run, so those keep the rebuild path."""

    def build(doc):
        doc.add_paragraph().add_run(f"见下表：{PLACEHOLDER}")

    rendered = _render(
        build,
        {
            KEY: [
                {
                    "type": "table",
                    "content": [
                        [{"type": "text", "content": "列1"}, {"type": "text", "content": "列2"}],
                        [{"type": "text", "content": "值1"}, {"type": "text", "content": "值2"}],
                    ],
                }
            ]
        },
    )

    assert len(rendered.tables) == 1
    assert rendered.tables[0].rows[0].cells[0].text == "列1"
    assert any("见下表：" in paragraph.text for paragraph in rendered.paragraphs)


def test_block_split_keeps_paragraph_style_and_run_formatting():
    """Splitting around a heading must not flatten the surviving text."""

    def build(doc):
        paragraph = doc.add_paragraph(style="Quote")
        paragraph.add_run("前缀 ").underline = True
        paragraph.add_run(PLACEHOLDER).underline = True
        paragraph.add_run(" 后缀____").underline = True

    rendered = _render(
        build,
        {
            KEY: [
                {"type": "text", "content": "正文"},
                {"type": "heading", "content": "小标题", "level": 3},
                {"type": "text", "content": "结尾"},
            ]
        },
    )
    before, heading, after = rendered.paragraphs[:3]

    assert before.style.name == "Quote"
    assert after.style.name == "Quote"
    assert heading.style.name == "Heading 3"
    # Text on both sides of the heading keeps the underline it had, and so does
    # the substituted value — it used to come out plain.
    assert all(run.underline is True for run in before.runs if run.text)
    assert all(run.underline is True for run in after.runs if run.text)
    assert before.text == "前缀 正文"
    assert after.text == "结尾 后缀____"


def test_block_split_keeps_each_run_its_own_format():
    """Mid-paragraph format changes used to be re-stamped with run[0]'s format."""

    def build(doc):
        paragraph = doc.add_paragraph()
        paragraph.add_run("普通 ")
        paragraph.add_run("加粗 ").bold = True
        paragraph.add_run(PLACEHOLDER).underline = True

    rendered = _render(
        build,
        {KEY: [{"type": "text", "content": "值"}, {"type": "heading", "content": "标题", "level": 2}]},
    )
    runs = [run for run in rendered.paragraphs[0].runs if run.text]

    assert [run.text for run in runs] == ["普通 ", "加粗 ", "值"]
    assert runs[0].bold is None
    assert runs[1].bold is True
    assert runs[2].underline is True
