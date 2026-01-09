import os
import tempfile
from pathlib import Path
from typing import IO, Dict, List, Any, Tuple
from urllib.parse import unquote, urlparse

import pandas as pd
import requests
from docx import Document
from docx.shared import Inches
from loguru import logger

from bisheng.core.storage.minio.minio_manager import get_minio_storage_sync
from bisheng.utils.util import _is_valid_url


def find_lcs(str1, str2):
    lstr1 = len(str1)
    lstr2 = len(str2)
    record = [[0 for i in range(lstr2 + 1)] for j in range(lstr1 + 1)]  # One more person
    maxNum = 0
    p = 0
    for i in range(lstr1):
        for j in range(lstr2):
            if str1[i] == str2[j]:
                record[i + 1][j + 1] = record[i][j] + 1
                if record[i + 1][j + 1] > maxNum:
                    maxNum = record[i + 1][j + 1]
                    p = i + 1

    return str1[p - maxNum: p], maxNum


class DocxTemplateRender(object):
    def __init__(self, filepath: str = None, file_content: IO[bytes] = None):
        self.filepath = filepath
        self.file_content = file_content
        if self.filepath:
            self.doc = Document(self.filepath)
        else:
            self.doc = Document(self.file_content)

    def _insert_image(self, paragraph, image_path: str, alt_text: str = "Images"):
        """
        Insert image in paragraph

        Args:
            paragraph: WordParagraph object
            image_path: Image file path
            alt_text: Set Alt Text(s)
        """
        logger.debug(f"[Simple Illustration] Start Inserting Pictures: {image_path}")
        try:
            if os.path.exists(image_path):
                # Check file size
                file_size = os.path.getsize(image_path)
                logger.debug(f"[Simple Illustration] Image file exists: {image_path}, size: {file_size}byte")

                # Insert an image with a maximum width of6Inch
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                logger.debug(f"[Simple Illustration] Ready atrunInsert a picture in therunQuantity: {len(paragraph.runs)}")

                run.add_picture(image_path, width=Inches(6))
                logger.info(f"[Simple Illustration] ✅ Successfully inserted image: {image_path}, size: {file_size}byte")
            else:
                # Image file does not exist, use original path
                logger.error(f"[Simple Illustration] ❌ Image file does not exist: {image_path}")
                if paragraph.runs:
                    paragraph.runs[0].text = image_path
                else:
                    paragraph.add_run(image_path)
        except Exception as e:
            # Inserting image failed, show originalURL
            logger.error(f"[Simple Illustration] ❌ Failed to insert image: {image_path}, Error type: {type(e).__name__}, Error-free: {str(e)}")
            if paragraph.runs:
                paragraph.runs[0].text = image_path
            else:
                paragraph.add_run(image_path)

    def _replace_placeholder_with_image(self, paragraph, placeholder, image_path, alt_text):
        """
        Precisely replace the placeholder in the paragraph with the picture, correctly handle the spanrunPlaceholder for

        Args:
            paragraph: WordParagraph object
            placeholder: Placeholder to replace
            image_path: Image file path
            alt_text: Image Alt Text
        """
        logger.debug(f"[Image Replacement] Start replacing placeholders: {placeholder} -> {image_path}")

        # Get the full text of the paragraph
        paragraph_text = paragraph.text

        # Find the location of the placeholder
        placeholder_start = paragraph_text.find(placeholder)
        if placeholder_start == -1:
            logger.warning(f"[Image Replacement] Placeholder not found: {placeholder}")
            return  # Placeholder does not exist

        placeholder_end = placeholder_start + len(placeholder)

        # Anchor placeholder inrunsPosition in
        current_pos = 0
        start_run_index = -1
        start_run_pos = 0
        end_run_index = -1
        end_run_pos = 0

        for i, run in enumerate(paragraph.runs):
            run_len = len(run.text)

            # Find the placeholder start position
            if start_run_index == -1 and current_pos + run_len > placeholder_start:
                start_run_index = i
                start_run_pos = placeholder_start - current_pos

            # Find the placeholder end position
            if current_pos + run_len >= placeholder_end:
                end_run_index = i
                end_run_pos = placeholder_end - current_pos
                break

            current_pos += run_len

        # Clear placeholder text
        if start_run_index == end_run_index:
            # Placeholder is the samerunand within
            run = paragraph.runs[start_run_index]
            run.text = run.text[:start_run_pos] + run.text[end_run_pos:]
        else:
            # Placeholders span multipleruns
            # Clear startrunSections in
            start_run = paragraph.runs[start_run_index]
            start_run.text = start_run.text[:start_run_pos]

            # Clear EndrunSections in
            end_run = paragraph.runs[end_run_index]
            end_run.text = end_run.text[end_run_pos:]

            # Clear middleruns
            for i in range(end_run_index - 1, start_run_index, -1):
                paragraph.runs[i].text = ""

        # Insert image at placeholder position
        # Find a suitable insertion location (first non-empty after cleaningrunafterwards
        insert_run = None
        for i in range(start_run_index, len(paragraph.runs)):
            if paragraph.runs[i].text or i == start_run_index:
                insert_run = paragraph.runs[i]
                break

        if insert_run is not None:
            # In therunInsert image after
            logger.debug(f"[Image Replacement] InsiderunPosition Insert Picture: {image_path}")
            self._insert_image_at_run(insert_run, image_path, alt_text)
            logger.info(f"[Image Replacement] Placeholder replacement complete: {placeholder} -> {image_path}")
        else:
            # If a suitable location is not found, use the original method
            logger.warning(f"[Image Replacement] No Fit Foundrunlocation, using alternate methods: {image_path}")
            self._insert_image(paragraph, image_path, alt_text)

    def _insert_image_at_run(self, run, image_path, alt_text):
        """
        In DesignationrunPosition Insert Picture

        Args:
            run: Word runObjects
            image_path: Image Path
            alt_text: Image Alt Text
        """
        try:
            from docx.shared import Inches

            # Checks to see if file exists.
            if not os.path.exists(image_path):
                logger.warning(f"[Image Rendering] Image file does not exist, use original path: {image_path}")
                run.text = image_path
                return

            # Check file size and format
            file_size = os.path.getsize(image_path)
            file_ext = os.path.splitext(image_path)[1].lower()
            logger.debug(f"[Image Rendering] Ready to insert image: path={image_path}, size={file_size}byte, ext={file_ext}")

            # Directly in the currentrunInsert image in
            run.add_picture(image_path, width=Inches(4))  # Default Width4Inch
            # Emptyruntext in to avoid displaying excess text
            run.text = ""
            logger.info(f"[Image Rendering] InsiderunLocation Insert Picture Successful: {image_path}, size={file_size}byte")
        except Exception as e:
            logger.error(f"[Image Rendering] runLocation Insert Picture Failed: {image_path}, Error-free: {e}")
            logger.debug(f"[Image Rendering] Error Details of error: {type(e).__name__}: {str(e)}")
            try:
                # Use original insertion method as backup
                logger.info(f"[Image Rendering] Try alternate insertion methods: {image_path}")
                # DapatkanrunParent Paragraph Object of
                paragraph = run._element.getparent()
                # Convert Topython-docxParagraph object
                from docx.text.paragraph import Paragraph

                para_obj = Paragraph(paragraph, run.part)
                self._insert_image(para_obj, image_path, alt_text)
                logger.info(f"[Image Rendering] Alternate method inserted successfully: {image_path}")
            except Exception as backup_e:
                logger.error(f"[Image Rendering] Alternate insertion methods also failed: {image_path}, Error-free: {backup_e}")
                logger.error(f"[Image Rendering] Alternate Method Detail Error: {type(backup_e).__name__}: {str(backup_e)}")
                # If all methods fail, use the original path
                run.text = image_path
                logger.warning(f"[Image Rendering] All insert methods failed, use original path text: {image_path}")

    def _replace_placeholder_in_structured_paragraph(self, paragraph, placeholder: str, table_data: List[List[str]]):
        """
        Simplified table replacement: insert the table directly at the placeholder without adding any structural tags
        """
        # Clear placeholder text
        self._clear_placeholder_from_paragraph(paragraph, placeholder)

        # Insert table directly at paragraph position
        self._insert_table(paragraph, table_data)

        logger.info(f"Form Replacement Complete of {len(table_data)} Row")

    def _clear_placeholder_from_paragraph(self, paragraph, placeholder):
        """
        Precisely clear placeholders from paragraphs to handle spanningrunCONDITION....&#x0D;

        Args:
            paragraph: WordParagraph object
            placeholder: Placeholders to be cleared
        """
        # Get the full text of the paragraph
        paragraph_text = paragraph.text

        # Find the location of the placeholder
        placeholder_start = paragraph_text.find(placeholder)
        if placeholder_start == -1:
            return  # Placeholder does not exist

        placeholder_end = placeholder_start + len(placeholder)

        # Anchor placeholder inrunsPosition in
        current_pos = 0
        start_run_index = -1
        start_run_pos = 0
        end_run_index = -1
        end_run_pos = 0

        for i, run in enumerate(paragraph.runs):
            run_len = len(run.text)

            # Find the placeholder start position
            if start_run_index == -1 and current_pos + run_len > placeholder_start:
                start_run_index = i
                start_run_pos = placeholder_start - current_pos

            # Find the placeholder end position
            if current_pos + run_len >= placeholder_end:
                end_run_index = i
                end_run_pos = placeholder_end - current_pos
                break

            current_pos += run_len

        # Clear placeholder text
        if start_run_index == end_run_index:
            # Placeholder is the samerunand within
            run = paragraph.runs[start_run_index]
            run.text = run.text[:start_run_pos] + run.text[end_run_pos:]
        else:
            # Placeholders span multipleruns
            # Clear startrunSections in
            start_run = paragraph.runs[start_run_index]
            start_run.text = start_run.text[:start_run_pos]

            # Clear EndrunSections in
            end_run = paragraph.runs[end_run_index]
            end_run.text = end_run.text[end_run_pos:]

            # Clear middleruns
            for i in range(end_run_index - 1, start_run_index, -1):
                paragraph.runs[i].text = ""

    def _process_resource_placeholders(self, doc, placeholder_map):
        """
        Work with hybrid placeholders in paragraphs in positional order

        Args:
            doc: WordDocument object  
            placeholder_map: Placeholder Mapping Dictionary
        """
        # Work with placeholders in all paragraphs
        paragraphs_to_process = list(doc.paragraphs)  # Create a copy as we may modify the paragraph structure

        for i, p in enumerate(paragraphs_to_process):
            paragraph_text = p.text
            if not paragraph_text:
                continue

            # Find all placeholders and their positions in the paragraph
            placeholders_with_positions = []
            for placeholder, resource_info in placeholder_map.items():
                pos = paragraph_text.find(placeholder)
                if pos != -1:
                    placeholders_with_positions.append({
                        'placeholder': placeholder,
                        'resource_info': resource_info,
                        'position': pos,
                        'end_position': pos + len(placeholder)
                    })

            if not placeholders_with_positions:
                continue

            # Sort by location from to post-processing
            placeholders_with_positions.sort(key=lambda x: x['position'])

            # Split paragraph into text segments and placeholder segments
            self._process_mixed_content_paragraph(doc, p, placeholders_with_positions, paragraph_text)

        # Working with Placeholders in Table Cells
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    for one in cell.paragraphs:
                        cell_text = one.text
                        if not cell_text:
                            continue

                        # Check for placeholders in cells
                        for placeholder, resource_info in placeholder_map.items():
                            if placeholder in cell_text:
                                if resource_info["type"] == "image":
                                    # Insert Actual Picture in Table Cell
                                    image_path = resource_info.get("local_path") or resource_info.get("path", "")
                                    if image_path and os.path.exists(image_path):
                                        try:
                                            # Insert a picture in a cell (this will empty the cell and insert the picture)
                                            self._insert_image_in_table_cell(cell, image_path)
                                            logger.info(f"✅ Picture successfully inserted in table cell: {image_path}")
                                            # Marker placeholder processed, no need to update text
                                            cell_text = ""
                                        except Exception as e:
                                            logger.error(f"❌ Table Cell Insert Picture Failed: {str(e)}")
                                            # Show file name on failure
                                            cell_text = cell_text.replace(placeholder, os.path.basename(image_path))
                                    else:
                                        # Image file does not exist, display path
                                        cell_text = cell_text.replace(placeholder,
                                                                      resource_info.get("path", placeholder))
                                elif resource_info["type"] == "excel":
                                    cell_text = cell_text.replace(placeholder, "[ExcelTable Filter]")
                                elif resource_info["type"] == "csv":
                                    cell_text = cell_text.replace(placeholder, "[CSVTable Filter]")
                                elif resource_info["type"] == "markdown_table":
                                    cell_text = cell_text.replace(placeholder, "[MarkdownTable Filter]")
                                logger.info(f"Process table cell placeholders: {placeholder}")

                        # Update cell text
                        if cell_text != one.text:
                            if one.runs:
                                one.runs[0].text = cell_text
                                for r_index in range(1, len(one.runs)):
                                    one.runs[r_index].text = ""
                            else:
                                one.add_run(cell_text)

    def _process_mixed_content_paragraph(self, doc, paragraph, placeholders_with_positions, original_text):
        """
        Working with paragraphs with mixed content - Insert images and tables inline
        
        Args:
            doc: WordDocument object
            paragraph: Original paragraph
            placeholders_with_positions: List of placeholder information by location
            original_text: Original paragraph text
        """
        # Extract style information from the original paragraph for subsequent paragraph creation
        original_style_info = self._extract_paragraph_style_info(paragraph)

        # Split the text into fragments, keeping the original order
        segments = []
        last_end = 0

        for item in placeholders_with_positions:
            # Add text before placeholder
            if item['position'] > last_end:
                text_before = original_text[last_end:item['position']]
                if text_before:  # Keep all text, including whitespaces
                    segments.append({
                        'type': 'text',
                        'content': text_before
                    })

            # Add the resource corresponding to the placeholder
            segments.append({
                'type': 'resource',
                'placeholder': item['placeholder'],
                'resource_info': item['resource_info']
            })

            last_end = item['end_position']

        # Add last remaining text
        if last_end < len(original_text):
            text_after = original_text[last_end:]
            if text_after:
                segments.append({
                    'type': 'text',
                    'content': text_after
                })

        # Clear original paragraph
        paragraph.clear()

        # Redesign: segmentation to keep document structure consistent
        current_paragraph = paragraph

        i = 0
        while i < len(segments):
            segment = segments[i]

            if segment['type'] == 'text':
                # Add text content directly without any cleanup
                current_paragraph.add_run(segment['content'])

            elif segment['type'] == 'resource':
                resource_info = segment['resource_info']

                if resource_info['type'] == 'image':
                    # Images can be really inlined in paragraphs
                    self._insert_inline_image(current_paragraph, resource_info['path'],
                                              resource_info.get('alt_text', ''))

                elif resource_info['type'] in ['excel', 'csv', 'markdown_table']:
                    # Table Inline Processing: Insert the table at its current position, then create a new paragraph for subsequent content
                    if resource_info['type'] == 'markdown_table':
                        table_data, alignments = self._markdown_table_to_data(resource_info["content"])
                    else:
                        table_data = resource_info.get('table_data', [["Table data parsing failed"]])
                        alignments = None

                    # Insert table immediately after current paragraph
                    table_element = self._create_table_element(table_data)

                    # Get the position of the current paragraph in the document
                    paragraph_element = current_paragraph._element
                    paragraph_parent = paragraph_element.getparent()
                    paragraph_index = list(paragraph_parent).index(paragraph_element)

                    # Insert table after current paragraph
                    paragraph_parent.insert(paragraph_index + 1, table_element)

                    # Create a new paragraph for the subsequent text and update itcurrent_paragraph
                    if i + 1 < len(segments):
                        next_paragraph = self._create_new_paragraph_after_table(paragraph_parent, paragraph_index + 1,
                                                                                original_style_info)
                        current_paragraph = next_paragraph

            i += 1

    def _insert_table_inline(self, current_paragraph, table_data, segments, current_index):
        """
        Insert tables inline for text continuity
        
        Args:
            current_paragraph: Current paragraph
            table_data: Form Data
            segments: All phrases
            current_index: Current Snippet Index
            
        Returns:
            int: Number of subsequent text fragments skipped
        """
        # 1. Insert table paragraph after current paragraph
        table_paragraph = self._create_new_paragraph_after(current_paragraph)
        self._insert_table(table_paragraph, table_data)

        # 2. Check if there is a subsequent text fragment and create a new one if there is one
        remaining_segments = segments[current_index + 1:]
        text_segments = [seg for seg in remaining_segments if seg['type'] == 'text']

        if text_segments:
            # Create a new paragraph for the subsequent text
            next_text_paragraph = self._create_new_paragraph_after(table_paragraph)

            # Add all subsequent text fragments to the new paragraph
            for j, seg in enumerate(text_segments):
                next_text_paragraph.add_run(seg['content'])

            # Returns the number of text fragments skipped (so that the main loop no longer processes these fragments)
            skipped_count = len(text_segments)
            return skipped_count
        else:
            return 0

    def _create_new_paragraph_after(self, paragraph):
        """
        Create a new paragraph after the specified paragraph
        
        Args:
            paragraph: Topic Paragraph reference
            
        Returns:
            Newly created paragraph object
        """
        from docx.oxml import parse_xml
        from docx.text.paragraph import Paragraph

        # Create a new paragraph element
        new_p_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:p>'
        new_p_element = parse_xml(new_p_xml)

        # Insert a new paragraph after the current paragraph
        paragraph._element.getparent().insert(
            list(paragraph._element.getparent()).index(paragraph._element) + 1,
            new_p_element
        )

        # Returns the wrapped paragraph object
        return Paragraph(new_p_element, paragraph._parent)

    def _insert_inline_image(self, paragraph, image_path: str, alt_text: str = ""):
        """
        Insert an image inline in the paragraph, and the image will be in the correct position in the text stream
        """
        try:
            # Use an existing image insertion method, but make sure it's inline
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(4.0))
            logger.info(f"[Inline Image] ✅ Successfully inserted image inline: {image_path}")
        except Exception as e:
            logger.error(f"[Inline Image] ❌ Failed to insert image: {str(e)}")
            # Insert original image path
            paragraph.add_run(image_path)

    def _insert_image_in_table_cell(self, cell, image_path: str):
        """
        Insert a picture in a table cell
        
        Args:
            cell: Table Cell Object
            image_path: Image file path
        """
        try:
            from docx.shared import Inches
            import os

            if not os.path.exists(image_path):
                logger.warning(f"[Table Cell Picture] Image file does not exist: {image_path}")
                return

            # Empty cell contents
            for paragraph in cell.paragraphs:
                paragraph.clear()

            # Insert image in first paragraph
            first_paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            run = first_paragraph.add_run()

            # Insert an image and set the appropriate size
            run.add_picture(image_path, width=Inches(2.0))  # Use smaller dimensions in table cells

            logger.info(f"[Table Cell Picture] ✅ Successfully inserted image: {image_path}")

        except Exception as e:
            logger.error(f"[Table Cell Picture] ❌ Failed to insert: {str(e)}")
            raise e

    def _csv_to_table(self, csv_path: str) -> List[List[str]]:
        """
        will beCSVConvert file to table data

        Args:
            csv_path: CSVFilePath

        Returns:
            List[List[str]]: Form Data
        """
        try:
            # read outCSVfiles, automatically detecting encodings and delimiters
            import csv
            import chardet

            # Test file code
            with open(csv_path, "rb") as f:
                raw_data = f.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result["encoding"] if encoding_result["encoding"] else "utf-8"

            table_data = []

            # Try a different delimiter
            delimiters = [",", ";", "\t", "|"]

            for delimiter in delimiters:
                try:
                    with open(csv_path, "r", encoding=encoding, newline="") as f:
                        # Read a small part first to detect the delimiter
                        sample = f.read(1024)
                        f.seek(0)

                        sniffer = csv.Sniffer()
                        try:
                            detected_delimiter = sniffer.sniff(sample).delimiter
                        except:
                            detected_delimiter = delimiter

                        reader = csv.reader(f, delimiter=detected_delimiter)
                        table_data = [row for row in reader]

                        # Delimiter is considered correct if data is read successfully and there are multiple columns
                        if table_data and len(table_data[0]) > 1:
                            break

                except Exception as e:
                    logger.debug(f"Try Delimiter '{delimiter}' Kalah: {str(e)}")
                    continue

            if not table_data:
                # If all delimiters fail, try a simple line-by-line read
                with open(csv_path, "r", encoding=encoding) as f:
                    lines = f.readlines()
                    table_data = [[line.strip()] for line in lines if line.strip()]

            # Cleanup data - Fully preserve the original table structure, including empty rows
            cleaned_data = []
            for row in table_data:
                cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                # Removes blank row filtering, leaving the original table structure intact
                cleaned_data.append(cleaned_row)

            logger.info(f"Successfully parsedCSVDoc.: {csv_path}, # of Lines: {len(cleaned_data)}")
            return cleaned_data

        except Exception as e:
            logger.error(f"analyzingCSVFile failed: {csv_path}, Error-free: {str(e)}")
            return [["CSVFile parsing failed", str(e)]]

    def _excel_to_table(self, excel_path: str) -> List[List[str]]:
        """
        will beExcelConvert files to tabular data, preserving data formats and types

        Args:
            excel_path: ExcelFilePath

        Returns:
            List[List[str]]: Form Data
        """

        try:
            # read outExcelfiles, keeping more of the original formatting
            df = pd.read_excel(excel_path, sheet_name=0, dtype=str, keep_default_na=False)

            # Limit table size to avoid overly large table effectsWordDocumentation
            max_rows = 500  # Max row count
            max_cols = 20  # Maximum Columns

            if len(df) > max_rows:
                logger.warning(f"ExcelToo many file lines({len(df)}), before interception{max_rows}Parade")
                df = df.head(max_rows)

            if len(df.columns) > max_cols:
                logger.warning(f"ExcelToo many file columns({len(df.columns)}), before interception{max_cols}column")
                df = df.iloc[:, :max_cols]

            # Convert to list format
            table_data = []

            # Process header, clean up column names
            headers = []
            for col in df.columns:
                col_str = str(col).strip()
                # <g id="Bold">Medical Treatment:</g>ExcelAutomatically generated column names (e.g.Unnamed: 0）
                if col_str.startswith("Unnamed:"):
                    col_str = ""  # Empty column name
                headers.append(col_str)
            table_data.append(headers)

            # Add data rows, smart formatting
            for _, row in df.iterrows():
                row_data = []
                for cell in row:
                    cell_str = str(cell).strip()

                    # Handle Empty Values
                    if cell_str.lower() in ["nan", "none", "null", ""]:
                        cell_str = ""

                    # Handle long numbers for readability
                    elif cell_str.replace(".", "").replace("-", "").isdigit():
                        try:
                            # Try formatting numbers
                            if "." in cell_str:
                                # Float, keep reasonable decimal places
                                num = float(cell_str)
                                if abs(num) >= 1000:
                                    cell_str = f"{num:,.2f}"  # Range Slider Thousand Separator
                                else:
                                    cell_str = f"{num:.2f}".rstrip("0").rstrip(".")
                            else:
                                # integer, adding thousands separator
                                num = int(cell_str)
                                if abs(num) >= 1000:
                                    cell_str = f"{num:,}"
                        except ValueError:
                            pass  # Keep original string

                    # Limit cell content length to avoid too long content affecting layout
                    if len(cell_str) > 100:
                        cell_str = cell_str[:97] + "..."

                    row_data.append(cell_str)

                table_data.append(row_data)

            # Make sure all rows have the same number of columns
            if table_data:
                max_cols = max(len(row) for row in table_data)
                for row in table_data:
                    while len(row) < max_cols:
                        row.append("")

            logger.info(f"Successfully parsedExcelDoc.: {excel_path}, size: {len(table_data)}Parade x {max_cols}column")
            return table_data

        except Exception as e:
            logger.error(f"analyzingExcelFile failed: {excel_path}, Error-free: {str(e)}")
            return [["ExcelFile parsing failed", str(e)]]

    def _markdown_table_to_data(self, markdown_table: str) -> Tuple[List[List[str]], List[str]]:
        """
        will beMarkdownTable to table data and parse alignment information

        Args:
            markdown_table: MarkdownTable Text

        Returns:
            tuple: (Form Data, Alignment Info List)
        """
        try:
            lines = [line.strip() for line in markdown_table.strip().split("\n") if line.strip()]

            if len(lines) < 2:
                logger.warning("MarkdownTable format is incomplete, at least table header and delimiter lines are required")
                return [["Format salah.", "The form is incomplete."]], ["left"]

            table_data = []
            alignments = []
            separator_found = False

            for i, line in enumerate(lines):
                # Check if it is a delimiter line
                if self._is_separator_line(line):
                    alignments = self._parse_alignments(line)
                    separator_found = True
                    continue

                # Parse Data Rows
                cells = self._parse_table_row(line)
                if cells:
                    # Clean cell contents
                    cleaned_cells = []
                    for cell in cells:
                        cleaned_cell = self._clean_cell_content(cell)
                        cleaned_cells.append(cleaned_cell)

                    table_data.append(cleaned_cells)

            # Validate table structure
            if not separator_found:
                logger.warning("MarkdownTable is missing a delimiter row, use default left alignment")
                alignments = ["left"] * (len(table_data[0]) if table_data else 1)

            # Make sure all rows have the same number of columns
            if table_data:
                max_cols = max(len(row) for row in table_data)
                for row in table_data:
                    while len(row) < max_cols:
                        row.append("")

                # Make sure the number of aligned messages matches the number of columns
                while len(alignments) < max_cols:
                    alignments.append("left")
                alignments = alignments[:max_cols]

            logger.info(f"Successfully parsedMarkdownTable, Size: {len(table_data)}Parade x {len(alignments)}column")
            logger.info(f"Column alignment: {alignments}")
            return table_data, alignments

        except Exception as e:
            logger.error(f"analyzingMarkdownTable failed: {str(e)}")
            return [["MarkdownTable parsing failed", str(e)]], ["left"]

    def _is_separator_line(self, line: str) -> bool:
        """
        Determine if yesMarkdownTable Separator Row

        Args:
            line: Table Row Content

        Returns:
            bool: Is Separator Row
        """
        # Remove leading and trailing|Symbols
        content = line.strip().strip("|").strip()
        if not content:
            return False

        # The delimiter line should mainly contain-And:characters. 
        cells = [cell.strip() for cell in content.split("|")]

        for cell in cells:
            if not cell:
                continue
            # Each cell should be predominantly composed of-And:Composition
            clean_cell = cell.replace("-", "").replace(":", "").strip()
            if clean_cell:  # Not a delimiter line if there are other characters
                return False

        return True

    def _parse_alignments(self, separator_line: str) -> List[str]:
        """
        Parse column alignment from delimiter rows

        Args:
            separator_line: Separator Row

        Returns:
            List[str]: Alignment List ("left", "center", "right")
        """
        alignments = []
        content = separator_line.strip().strip("|").strip()
        cells = [cell.strip() for cell in content.split("|")]

        for cell in cells:
            if not cell:
                alignments.append("left")
                continue

            # Determine alignment
            if cell.startswith(":") and cell.endswith(":"):
                alignments.append("center")  # :---:
            elif cell.endswith(":"):
                alignments.append("right")  # ---:
            else:
                alignments.append("left")  # --- OR :---

        return alignments

    def _parse_table_row(self, line: str) -> List[str]:
        """
        Parse table rows, handle escapes and special characters

        Args:
            line: Table Row Content

        Returns:
            List[str]: Cell Contents List
        """
        # Remove leading and trailing|Symbols
        content = line.strip()
        if content.startswith("|"):
            content = content[1:]
        if content.endswith("|"):
            content = content[:-1]

        # Split cells, but handle escapes|
        cells = []
        current_cell = ""
        escaped = False

        for char in content:
            if escaped:
                current_cell += char
                escaped = False
            elif char == "\\":
                escaped = True
                current_cell += char
            elif char == "|":
                cells.append(current_cell.strip())
                current_cell = ""
            else:
                current_cell += char

        # Add last cell
        if current_cell or cells:  # Process Empty Rows
            cells.append(current_cell.strip())

        return cells

    def _clean_cell_content(self, cell: str) -> str:
        """
        Clean up cell contents, processMarkdownFormat

        Args:
            cell: Raw Cell Contents

        Returns:
            str: Post-Cleanup Content
        """
        if not cell:
            return ""

        # Remove extra spaces
        cleaned = cell.strip()

        # <g id="Bold">Medical Treatment:</g>MarkdownFormatting tags (simplified processing)
        # Remove bold mark
        cleaned = cleaned.replace("**", "")
        # Remove italic mark
        cleaned = cleaned.replace("*", "")
        # Remove code marker
        cleaned = cleaned.replace("`", "")

        # Working with link formats [text](url) -> text
        import re

        cleaned = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", cleaned)

        # Limit length
        if len(cleaned) > 100:
            cleaned = cleaned[:97] + "..."

        return cleaned

    def _insert_markdown_table(self, paragraph, table_data: List[List[str]], alignments: List[str]):
        """
        Insert ExclusivelyMarkdownTable with support for alignment information

        Args:
            paragraph: WordParagraph object
            table_data: Form Data
            alignments: Alignment Info List
        """
        try:
            if not table_data:
                paragraph.add_run("Table data is empty")
                return

            rows = len(table_data)
            cols = len(alignments)

            # Simplify processing: insert a table directly at the paragraph position
            paragraph_element = paragraph._element
            paragraph_parent = paragraph_element.getparent()

            # Create Table
            table = self.doc.add_table(rows=rows, cols=cols)
            table_element = table._tbl

            # Insert table directly after paragraph
            paragraph_index = list(paragraph_parent).index(paragraph_element)
            paragraph_parent.insert(paragraph_index + 1, table_element)

            # PengaturanMarkdownTable-specific styles
            try:
                table.style = "Light List - Accent 1"  # Refreshing list style
            except Exception:
                try:
                    table.style = "Table Grid"  # Alternative style
                except Exception:
                    pass

            # Populate table data and set alignment
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    if j < cols:
                        cell = table.cell(i, j)
                        cell.text = str(cell_data)

                        # accordingMarkdownAlignment Info Set Cell Alignment
                        cell_paragraphs = cell.paragraphs
                        if cell_paragraphs and j < len(alignments):
                            alignment = alignments[j]
                            if alignment == "center":
                                cell_paragraphs[0].alignment = 1  # Center
                            elif alignment == "right":
                                cell_paragraphs[0].alignment = 2  # Halign Right
                            else:  # left
                                cell_paragraphs[0].alignment = 0  # Align left

            # MarkdownTable style: The first row is usually the header
            if rows > 0:
                for cell in table.rows[0].cells:
                    # Header Style
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                    # Set header background color (light gray, better forMarkdownSTYLE #
                    try:
                        from docx.oxml import parse_xml

                        shading_elm = parse_xml(
                            r'<w:shd {} w:fill="F2F2F2"/>'.format(
                                'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                            )
                        )
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                    except Exception:
                        pass

            # Set Table Layout
            try:
                table.autofit = True
                from docx.shared import Inches

                table.width = Inches(6.5)
            except Exception:
                pass

            # Set a more compact line height (MarkdownSTYLE #
            try:
                for row in table.rows:
                    row.height = Inches(0.25)  # .ExcelForms are more compact
                    for cell in row.cells:
                        # Smaller margin
                        cell.margin_left = Inches(0.03)
                        cell.margin_right = Inches(0.03)
                        cell.margin_top = Inches(0.01)
                        cell.margin_bottom = Inches(0.01)
            except Exception:
                pass

            logger.info(f"Succesfully insertedMarkdownTable, Size: {rows}x{cols}Align: {alignments}")

        except Exception as e:
            paragraph.add_run(f"Table insertion failed: {str(e)}")
            logger.error(f"InsertMarkdownTable failed: {str(e)}")

    def _insert_table_at_position(self, paragraph, table_data: List[List[str]]):
        """
        Insert a table at the current position of the paragraph, and the table will appear before the paragraph

        Args:
            paragraph: WordParagraph object
            table_data: Form Data
        """
        try:
            if not table_data:
                paragraph.add_run("[Table data is empty]")
                return

            rows = len(table_data)
            cols = max(len(row) for row in table_data) if table_data else 1

            # Get the position of the paragraph in the document
            paragraph_element = paragraph._element
            paragraph_parent = paragraph_element.getparent()
            paragraph_index = list(paragraph_parent).index(paragraph_element)

            # Create Table
            table = self.doc.add_table(rows=rows, cols=cols)
            table_element = table._tbl

            # Insert table before paragraph (so the table appears in the current content position)
            paragraph_parent.insert(paragraph_index, table_element)

            # Fill in the table data and style it
            self._fill_and_style_table(table, table_data)

        except Exception as e:
            paragraph.add_run(f"[Table insertion failed: {str(e)}]")
            logger.error(f"Failed to insert table: {str(e)}")

    def _insert_table(self, paragraph, table_data: List[List[str]]):
        """
        Insert high quality table after paragraph

        Args:
            paragraph: WordParagraph object
            table_data: Form Data
        """
        try:
            if not table_data:
                paragraph.add_run("[Table data is empty]")
                return

            # Add table after paragraph
            rows = len(table_data)
            cols = max(len(row) for row in table_data) if table_data else 1

            # Simplify processing: insert a table directly at the paragraph position
            paragraph_element = paragraph._element
            paragraph_parent = paragraph_element.getparent()

            # Create Table
            table = self.doc.add_table(rows=rows, cols=cols)
            table_element = table._tbl

            # Insert table directly after paragraph
            paragraph_index = list(paragraph_parent).index(paragraph_element)
            paragraph_parent.insert(paragraph_index + 1, table_element)

            # Fill in the table data and style it
            self._fill_and_style_table(table, table_data)

        except Exception as e:
            paragraph.add_run(f"[Table insertion failed: {str(e)}]")
            logger.error(f"Failed to insert table: {str(e)}")

    def _fill_and_style_table(self, table, table_data: List[List[str]]):
        """
        Fill in the table data and style it
        
        Args:
            table: WordTable Object
            table_data: Form Data
        """
        rows = len(table_data)
        # Calculate the maximum number of columns, taking into account possible empty rows
        cols = max((len(row) for row in table_data if row), default=1)

        # Set a more professional table style
        try:
            # Try a better built-in style
            table.style = "Light Shading - Accent 1"  # Light Shadow Style
        except Exception:
            try:
                table.style = "Table Grid"  # Alternative style
            except Exception:
                pass  # If the style does not exist, use the default style

        # Populate table data
        for i, row_data in enumerate(table_data):
            # Ensure that empty lines can also be handled correctly
            current_row = row_data if row_data else []

            # Populate all columns of the row
            for j in range(cols):
                cell = table.cell(i, j)
                if j < len(current_row):
                    cell_data = current_row[j]
                else:
                    cell_data = ""  # Empty Cells

                cell.text = str(cell_data)

                # Set cell alignment
                cell_paragraphs = cell.paragraphs
                if cell_paragraphs:
                    # Numbers aligned right, text aligned left
                    cell_text = str(cell_data).strip()
                    if self._is_number(cell_text):
                        cell_paragraphs[0].alignment = 2  # Halign Right
                    else:
                        cell_paragraphs[0].alignment = 0  # Align left

        # Set header style (first line)
        if rows > 0:
            for cell in table.rows[0].cells:
                # Align Header Centered
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1  # Center Alignment
                    for run in paragraph.runs:
                        run.bold = True

                # Try setting the header background color (if supported)
                try:
                    from docx.oxml import parse_xml

                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="D9E2F3"/>'.format(
                            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                        )
                    )
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                except Exception:
                    pass  # If setting the background color fails, continue

        # Auto Fit Column Width
        try:
            table.autofit = True
            # Set table width to page width
            from docx.shared import Inches

            table.width = Inches(6.5)  # ca.A4Width (%)
        except Exception:
            pass

        # Set row height and cell margin
        try:
            for row in table.rows:
                row.height = Inches(0.3)  # Set Row Heights
                for cell in row.cells:
                    # Set cell margin
                    cell.margin_left = Inches(0.05)
                    cell.margin_right = Inches(0.05)
                    cell.margin_top = Inches(0.02)
                    cell.margin_bottom = Inches(0.02)
        except Exception:
            pass

        logger.info(f"Successfully set table style and data, size: {rows}x{cols}")

    def _create_table_element(self, table_data: List[List[str]]):
        """
        Create a table element
        
        Args:
            table_data: Form Data [[row1_col1, row1_col2], [row2_col1, row2_col2]]
        
        Returns:
            Form Elements
        """
        try:
            # Handling completely data-free scenarios: creating minimal tables
            if not table_data:
                logger.info("Table data isNoneor an empty list to create a minimal table:1x1")
                table = self.doc.add_table(rows=1, cols=1)
                table.cell(0, 0).text = ""  # Empty Cells
                return table._tbl

            # Calculate table size: keep original structure even if all rows are empty
            rows = len(table_data)
            if rows == 0:
                logger.info("The number of rows in the table is0, create a minimal table:1x1")
                table = self.doc.add_table(rows=1, cols=1)
                table.cell(0, 0).text = ""
                return table._tbl

            # Calculate the maximum number of columns, if all rows are empty, the default is1column
            cols = max((len(row) for row in table_data if row), default=1)

            logger.info(f"Create Table{rows}Parade x {cols}column")
            table = self.doc.add_table(rows=rows, cols=cols)

            # Fill in the table data and style it
            self._fill_and_style_table(table, table_data)

            return table._tbl

        except Exception as e:
            logger.error(f"Form element creation failed: {str(e)}", exc_info=True)
            return None

    def _create_new_paragraph_after_table(self, parent, table_index, style_info=None):
        """
        Create new paragraph after table
        
        Args:
            parent: Parent Container
            table_index: Index of the table in the parent container
            style_info: Style information dictionary for applying to new paragraphs
        
        Returns:
            New Paragraph Object
        """
        try:
            # Create a new paragraph
            new_paragraph = self.doc.add_paragraph()

            # Apply style information
            if style_info:
                self._apply_style_info(new_paragraph, style_info)
                alignment_info = f"inherit styles, aligning={style_info.get('alignment', 'None')}"
            else:
                # Back of pocket: set to left alignment
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                alignment_info = "Default left alignment"

            paragraph_element = new_paragraph._element

            # Insert new paragraph after table
            parent.insert(table_index + 1, paragraph_element)

            return new_paragraph

        except Exception as e:
            logger.error(f"Failed to create new paragraph: {str(e)}", exc_info=True)
            return None

    def _extract_paragraph_style_info(self, paragraph):
        """
        Extract full style information for paragraphs
        
        Args:
            paragraph: Paragraph object
            
        Returns:
            dict: Dictionary containing paragraph style information
        """
        try:
            style_info = {
                'alignment': paragraph.alignment,
                'paragraph_format': {},
                'style_name': None
            }

            # Extract paragraph formatting information
            if hasattr(paragraph, 'paragraph_format'):
                pf = paragraph.paragraph_format
                style_info['paragraph_format'] = {
                    'space_before': pf.space_before,
                    'space_after': pf.space_after,
                    'line_spacing': pf.line_spacing,
                    'left_indent': pf.left_indent,
                    'right_indent': pf.right_indent,
                    'first_line_indent': pf.first_line_indent,
                }

            # Extract Style Name
            if hasattr(paragraph, 'style') and paragraph.style:
                style_info['style_name'] = paragraph.style.name

            return style_info

        except Exception as e:
            logger.warning(f"Failed to extract paragraph style: {str(e)}")
            return {'alignment': None, 'paragraph_format': {}, 'style_name': None}

    def _apply_style_info(self, paragraph, style_info):
        """
        Apply style information to paragraphs
        
        Args:
            paragraph: Target Paragraph Object
            style_info: Style Information Dictionary
        """
        try:
            # Apply alignment
            if style_info.get('alignment') is not None:
                paragraph.alignment = style_info['alignment']

            # Apply Paragraph Formatting
            paragraph_format = style_info.get('paragraph_format', {})
            if paragraph_format:
                pf = paragraph.paragraph_format

                for attr_name, value in paragraph_format.items():
                    if value is not None and hasattr(pf, attr_name):
                        try:
                            setattr(pf, attr_name, value)
                        except Exception:
                            pass  # Ignore formatting application failed

            # Apply style name
            style_name = style_info.get('style_name')
            if style_name:
                try:
                    paragraph.style = style_name
                except Exception:
                    pass  # Ignore style application failed

        except Exception as e:
            logger.warning(f"Failed to apply style information: {str(e)}")

    def _is_number(self, text: str) -> bool:
        """
        Determine if the text is a number

        Args:
            text: Text to judge

        Returns:
            bool: Is numeric
        """
        if not text:
            return False

        # Remove thousands separator and spaces
        clean_text = text.replace(",", "").replace(" ", "").replace("%", "")

        try:
            float(clean_text)
            return True
        except ValueError:
            return False

    def render(self, template_def, resources: Dict[str, List[Dict[str, Any]]] = None):
        """
        Render templates for image and table insertion

        Args:
            template_def: Template Definition List
            resources: A dictionary of resource information, including pictures,Exceldocumentation   andMarkdownTable Filter
        """
        doc = self.doc

        # If no incomingresources, using an empty dictionary
        if resources is None:
            resources = {"images": [], "excel_files": [], "csv_files": [], "markdown_tables": []}

        # Create placeholder-to-resource mapping
        placeholder_map = {}

        # Image placeholder mapping
        for img_info in resources.get("images", []):
            placeholder_map[img_info["placeholder"]] = {
                "type": "image",
                "path": img_info["local_path"],
                "alt_text": img_info["alt_text"],
                "resource_type": img_info["type"],
            }

        # ExcelFile placeholder mapping
        for excel_info in resources.get("excel_files", []):
            placeholder_map[excel_info["placeholder"]] = {
                "type": "excel",
                "path": excel_info["local_path"],
                "resource_type": excel_info["type"],
                "table_data": excel_info.get("table_data", []),  # Tambahtable_dataData field
                "alignments": excel_info.get("alignments", None),  # TambahalignmentsData field
            }

        # CSVFile placeholder mapping
        for csv_info in resources.get("csv_files", []):
            placeholder_map[csv_info["placeholder"]] = {
                "type": "csv",
                "file_name": csv_info.get("file_name", "UnknownCSVDoc."),
                "table_data": csv_info.get("table_data", []),
                "resource_type": csv_info.get("type", "content"),
            }

        # MarkdownTable placeholder mapping
        for table_info in resources.get("markdown_tables", []):
            placeholder_map[table_info["placeholder"]] = {"type": "markdown_table", "content": table_info["content"]}

        # Original Text Replacement Logic
        for replace_info in template_def:
            k1 = replace_info[0]
            v1 = replace_info[1]

            # Work with placeholders in tables
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if k1 in cell.text:
                            for one in cell.paragraphs:
                                if k1 in one.text:
                                    # Checks for placeholders requiring special handling
                                    cell_text = one.text.replace(k1, v1)

                                    # Processing Placeholders
                                    for placeholder, resource_info in placeholder_map.items():
                                        if placeholder in cell_text:
                                            if resource_info["type"] == "image":
                                                # Insert Actual Picture in Table Cell
                                                image_path = resource_info.get("local_path") or resource_info.get(
                                                    "path", "")
                                                if image_path and os.path.exists(image_path):
                                                    try:
                                                        # Clear cell text
                                                        cell_text = cell_text.replace(placeholder, "")
                                                        # Insert Picture in Cell
                                                        self._insert_image_in_table_cell(cell, image_path)
                                                        logger.info(f"✅ Picture successfully inserted in table cell: {image_path}")
                                                    except Exception as e:
                                                        logger.error(f"❌ Table Cell Insert Picture Failed: {str(e)}")
                                                        # Show file name on failure
                                                        cell_text = cell_text.replace(placeholder,
                                                                                      os.path.basename(image_path))
                                                else:
                                                    # Image file does not exist, display path
                                                    cell_text = cell_text.replace(placeholder, resource_info.get("path",
                                                                                                                 placeholder))
                                            elif resource_info["type"] == "excel":
                                                cell_text = cell_text.replace(placeholder, "[ExcelTable Filter]")
                                            elif resource_info["type"] == "csv":
                                                cell_text = cell_text.replace(placeholder, "[CSVTable Filter]")
                                            elif resource_info["type"] == "markdown_table":
                                                cell_text = cell_text.replace(placeholder, "[MarkdownTable Filter]")

                                    one.runs[0].text = cell_text
                                    for r_index, r in enumerate(one.runs):
                                        if r_index == 0:
                                            continue
                                        r.text = ""

            # Processing Placeholders in Paragraphs
            for p in doc.paragraphs:
                if k1 not in p.text:
                    continue

                runs_cnt = len(p.runs)
                s_e = []
                i = 0
                while i < runs_cnt:
                    new_i = i + 1
                    for j in range(i + 1, runs_cnt + 1):
                        part_text = "".join([r.text for r in p.runs[i:j]])
                        if k1 in part_text:
                            # Find the smallest range containingk1right of privacyruns
                            tmp_i, tmp_j = i, j
                            while tmp_i <= tmp_j:
                                tmp_part_text = "".join([r.text for r in p.runs[tmp_i:tmp_j]])
                                if k1 in tmp_part_text:
                                    tmp_i += 1
                                    continue
                                else:
                                    tmp_i -= 1
                                    break
                            s_e.append((tmp_i, j))
                            new_i = j
                            break
                    i = new_i

                for one in s_e:
                    s, e = one
                    assert e > 0, [r.text for r in p.runs]

                    if e - s == 1:
                        replace_mapping = [(k1, v1)]
                    elif e - s == 2:
                        s_tgt_text = p.runs[s].text
                        comm_str, max_num = find_lcs(k1, s_tgt_text)
                        assert k1.startswith(comm_str)
                        p1 = comm_str
                        p2 = k1[max_num:]
                        n = len(v1)
                        sub_n1 = int(1.0 * len(p1) / (len(p1) + len(p2)) * n)
                        replace_mapping = [(p1, v1[:sub_n1]), (p2, v1[sub_n1:])]
                    elif e - s == 3:
                        m_text = p.runs[s + 1].text
                        head_tail = k1.split(m_text, 1)
                        assert len(head_tail) == 2
                        h_text = head_tail[0]
                        t_text = head_tail[1]
                        replace_mapping = [(h_text, ""), (m_text, v1), (t_text, "")]
                    else:
                        m_texts = [p.runs[i].text for i in range(s + 1, e - 1)]
                        m_text = "".join(m_texts)
                        head_tail = k1.split(m_text, 1)
                        assert len(head_tail) == 2
                        h_text = head_tail[0]
                        t_text = head_tail[1]
                        replace_mapping = [(h_text, "")]
                        replace_mapping.append((m_texts[0], v1))
                        for text in m_texts[1:]:
                            replace_mapping.append((text, ""))
                        replace_mapping.append((t_text, ""))

                    for i in range(s, e):
                        _k, _v = replace_mapping[i - s]
                        p.runs[i].text = p.runs[i].text.replace(_k, _v)

        # Unify resource placeholders after all variable replacements are complete
        self._process_resource_placeholders(doc, placeholder_map)

        # Add Final Document Content Check
        self._log_final_document_content(doc)

        return doc

    def _log_final_document_content(self, doc):
        """Review final document content"""
        try:
            # Check for unprocessedfile contentlabel
            for i, paragraph in enumerate(doc.paragraphs):
                paragraph_text = paragraph.text.strip()
                if paragraph_text and 'file content' in paragraph_text.lower():
                    logger.warning(f"Unprocessed foundfile contentLabel, paragraph{i}: {paragraph_text[:200]}...")
        except Exception as e:
            logger.error(f"Failed to check document content: {str(e)}")


def test_replace_string(template_file, kv_dict: dict, file_name: str):
    # If the file is a web path, download it to a temporary file, and use that
    if not os.path.isfile(template_file) and _is_valid_url(template_file):
        r = requests.get(template_file)

        if r.status_code != 200:
            raise ValueError("Check the url of your file; returned status code %s" % r.status_code)

        temp_dir = tempfile.TemporaryDirectory()
        temp_file = Path(temp_dir.name) / unquote(urlparse(template_file).path.split("/")[-1])
        with open(temp_file, mode="wb") as f:
            f.write(r.content)

        template_file = temp_file
    elif not os.path.isfile(template_file):
        raise ValueError("File path %s is not a valid file or url" % template_file)

    template_dict = []
    for k, v in kv_dict.items():
        template_dict.append(["{{" + k + "}}", v])

    doc = DocxTemplateRender(str(template_file))
    output = doc.render(template_dict)

    temp_dir = tempfile.TemporaryDirectory()
    temp_file = Path(temp_dir.name) / file_name
    output.save(temp_file)
    minio_client = get_minio_storage_sync()
    minio_client.put_object_sync(bucket_name=minio_client.bucket, object_name=file_name, file=temp_file)

    return file_name
