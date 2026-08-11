#!/usr/bin/env python3
"""Read out and health-check a .docx, for the BiSheng code interpreter.

    python skills/bisheng-docx/scripts/inspect_docx.py output/x.docx
    python skills/bisheng-docx/scripts/inspect_docx.py output/x.docx --content-only
    python skills/bisheng-docx/scripts/inspect_docx.py output/x.docx --check-only

Replaces two things the official skill relies on and this environment lacks:
``pandoc -t markdown`` for reading (pandoc exists in the release image but not
on every hand-built host, and it drops the structural detail needed to plan an
edit), and "render to JPG and look at every page", which assumes a model that
can see images.

The single highest-value check is the missing ``w:eastAsia`` font: it is
invisible in python-docx, invisible in extracted text, and shows up only when
the user opens the file and finds every Chinese character in the wrong face.

Always exits 0 — the executor discards stdout on a non-zero exit.
"""

import argparse
import os
import sys
import traceback

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:  # pragma: no cover
    print("[FATAL] python-docx 不可用；无法体检。")
    sys.exit(0)

PLACEHOLDERS = ["待填", "待补", "TODO", "XXX", "占位", "请填写", "TBD", "lorem ipsum", "示例文字"]
CJK_RANGE = ("一", "鿿")

findings: list[tuple[str, str]] = []


def add(level: str, message: str) -> None:
    findings.append((level, message))


def has_cjk(text: str) -> bool:
    return any(CJK_RANGE[0] <= ch <= CJK_RANGE[1] for ch in text)


def run_east_asia(run) -> str | None:
    """The w:eastAsia face actually written for this run, if any."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn("w:eastAsia"))


def style_east_asia(style) -> str | None:
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn("w:eastAsia"))


def iter_block_paragraphs(doc):
    """Body paragraphs plus every paragraph inside a table cell."""
    for paragraph in doc.paragraphs:
        yield paragraph, None
    for t_index, table in enumerate(doc.tables):
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    yield paragraph, f"表{t_index + 1}[{r_index + 1},{c_index + 1}]"


# --------------------------------------------------------------------------- #
# Content read-out
# --------------------------------------------------------------------------- #


def dump_content(path: str, max_paras: int) -> None:
    doc = Document(path)
    section = doc.sections[0]
    width_cm = section.page_width / 360000
    height_cm = section.page_height / 360000

    print("=" * 72)
    print(f"内容 · {os.path.basename(path)}")
    print("=" * 72)
    print(
        f"页面 {width_cm:.1f}×{height_cm:.1f}cm  "
        f"页边距 上{section.top_margin / 360000:.1f} 下{section.bottom_margin / 360000:.1f} "
        f"左{section.left_margin / 360000:.1f} 右{section.right_margin / 360000:.1f}cm  "
        f"段落 {len(doc.paragraphs)}  表格 {len(doc.tables)}  图片 {len(doc.inline_shapes)}"
    )
    print()

    shown = 0
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if shown >= max_paras:
            print("   … 其余段落未显示（--max-paras 可调）")
            break
        style = paragraph.style.name if paragraph.style is not None else "?"
        tag = f"[{style}]" if style != "Normal" else ""
        print(f"P{index:<4} {tag} {text if len(text) <= 110 else text[:107] + '…'}")
        shown += 1

    for t_index, table in enumerate(doc.tables):
        print(f"\n## 表 {t_index + 1} · {len(table.rows)} 行 × {len(table.columns)} 列")
        for r_index, row in enumerate(table.rows[:8]):
            cells = [c.text.strip().replace("\n", " ")[:24] for c in row.cells]
            print("   | " + " | ".join(cells) + " |")
        if len(table.rows) > 8:
            print(f"   … 其余 {len(table.rows) - 8} 行未显示")


# --------------------------------------------------------------------------- #
# Health checks
# --------------------------------------------------------------------------- #


def check_fonts(doc) -> None:
    normal_ea = style_east_asia(doc.styles["Normal"])
    if not normal_ea:
        add(
            "WARN",
            "Normal 样式没有设置 w:eastAsia 中文字体 —— 中文会退回主题字体（通常是等线/宋体），"
            "和你指定的西文字体对不上。用 apply_chinese_defaults(doc) 一次性设好。",
        )

    bad_runs: list[str] = []
    checked = 0
    for paragraph, where in iter_block_paragraphs(doc):
        for run in paragraph.runs:
            if not has_cjk(run.text):
                continue
            checked += 1
            latin = run.font.name
            east = run_east_asia(run)
            # Only a run that names a Latin face but no CJK face is broken; a run
            # inheriting everything from a correctly-set style is fine.
            if latin and not east:
                label = where or f"「{run.text.strip()[:16]}」"
                if label not in bad_runs:
                    bad_runs.append(label)
    if bad_runs:
        add(
            "ERROR",
            f"{len(bad_runs)} 处中文 run 只设了西文字体、没设 w:eastAsia："
            f"{', '.join(bad_runs[:6])}{' …' if len(bad_runs) > 6 else ''}。"
            f"`run.font.name = '微软雅黑'` 对中文无效，必须用 set_run_font()。",
        )
    if checked == 0:
        add("INFO", "文档里没有中文 run（纯英文文档？）")


def check_headings(doc) -> None:
    levels = []
    for paragraph in doc.paragraphs:
        name = paragraph.style.name if paragraph.style is not None else ""
        if name.startswith("Heading ") and paragraph.text.strip():
            try:
                levels.append((int(name.split()[-1]), paragraph.text.strip()))
            except ValueError:
                continue
    if not levels:
        add("WARN", "文档里没有用内置 Heading 样式的标题 —— 目录生成不出来，导航窗格也是空的")
        return
    previous = 0
    for level, text in levels:
        if previous and level > previous + 1:
            add("WARN", f"标题层级从 H{previous} 直接跳到 H{level}：「{text[:24]}」")
        previous = level


def check_tables(doc, section) -> None:
    usable_cm = (section.page_width - section.left_margin - section.right_margin) / 360000
    for index, table in enumerate(doc.tables):
        if not table.rows:
            add("WARN", f"表 {index + 1} 是空的")
            continue
        widths = [c.width for c in table.rows[0].cells]
        if all(w is not None for w in widths):
            total_cm = sum(w / 360000 for w in widths)
            if total_cm > usable_cm + 0.2:
                add(
                    "ERROR",
                    f"表 {index + 1} 总宽 {total_cm:.1f}cm 超过版心 {usable_cm:.1f}cm → 右侧列会被截出页面",
                )
        elif any(w is None for w in widths):
            add(
                "INFO",
                f"表 {index + 1} 没有设列宽，Word 会自动分配；列多时中文列容易被压成竖排。"
                f"建议给 add_table 传 widths_cm。",
            )
        for r_index, row in enumerate(table.rows):
            if len(row.cells) != len(table.columns):
                add("WARN", f"表 {index + 1} 第 {r_index + 1} 行单元格数与列数不一致（可能有合并）")
                break


def check_images(doc, section) -> None:
    usable_cm = (section.page_width - section.left_margin - section.right_margin) / 360000
    for index, shape in enumerate(doc.inline_shapes):
        width_cm = shape.width / 360000
        if width_cm > usable_cm + 0.2:
            add("ERROR", f"图 {index + 1} 宽 {width_cm:.1f}cm 超过版心 {usable_cm:.1f}cm → 会溢出页面")


def check_text(doc) -> None:
    empty_streak = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            empty_streak += 1
            if empty_streak == 3:
                add("INFO", "连续 3 个以上空段落 —— 用段前段后间距代替空行")
        else:
            empty_streak = 0
        for token in PLACEHOLDERS:
            if token in text:
                add("WARN", f"残留占位文本「{token}」：{text[:40]}")
                break
        if "\n" in paragraph.text:
            add("WARN", f"段落里有换行符而不是独立段落：{text[:40]} —— 每段用一个 Paragraph")


def check_fields(doc) -> None:
    xml = doc.element.xml
    has_toc = "TOC \\o" in xml or "TOC \\O" in xml
    if has_toc:
        settings_xml = doc.settings.element.xml
        if "updateFields" not in settings_xml:
            add(
                "ERROR",
                "有目录域但没开 updateFields —— 用户打开时目录仍是占位文字，看起来像坏了。"
                "调用 enable_update_fields(doc)（add_toc 已内置）。",
            )
    if "PAGE " not in xml and len(doc.paragraphs) > 40:
        add("INFO", "文档较长但页脚没有页码域")


def run_checks(path: str) -> None:
    doc = Document(path)
    section = doc.sections[0]
    check_fonts(doc)
    check_headings(doc)
    check_tables(doc, section)
    check_images(doc, section)
    check_text(doc)
    check_fields(doc)

    print()
    print("=" * 72)
    print("体检")
    print("=" * 72)

    order = {"ERROR": 0, "WARN": 1, "INFO": 2}
    counts = {"ERROR": 0, "WARN": 0, "INFO": 0}
    for level, _ in findings:
        counts[level] += 1
    for level, message in sorted(findings, key=lambda f: order[f[0]]):
        print(f"[{level}] {message}")
    if not findings:
        print("（无发现）")

    print()
    print(f"ERROR {counts['ERROR']} · WARN {counts['WARN']} · INFO {counts['INFO']}")
    if counts["ERROR"]:
        print("结论: 不通过 —— ERROR 必须全部修完再交付。")
    elif counts["WARN"]:
        print("结论: 有条件通过 —— WARN 逐条复核，确认无误可交付。")
    else:
        print("结论: 通过。")


def main() -> None:
    parser = argparse.ArgumentParser(description="读出并体检一个 .docx")
    parser.add_argument("path")
    parser.add_argument("--content-only", action="store_true")
    parser.add_argument("--check-only", action="store_true")
    parser.add_argument("--max-paras", type=int, default=80)
    args = parser.parse_args()

    if not os.path.exists(args.path):
        print(f"[FATAL] 文件不存在: {args.path}")
        print("提示：代码执行器 cwd 就是工作区根，用相对路径 `output/x.docx`，不要写 `/output/x.docx`。")
        return

    try:
        if not args.check_only:
            dump_content(args.path, args.max_paras)
        if not args.content_only:
            run_checks(args.path)
    except Exception:
        print("[FATAL] 体检过程本身出错：")
        traceback.print_exc(file=sys.stdout)


if __name__ == "__main__":
    main()
    sys.exit(0)
