from typing import IO

from docx import Document


def find_lcs(str1, str2):
    lstr1 = len(str1)
    lstr2 = len(str2)
    record = [[0 for i in range(lstr2 + 1)] for j in range(lstr1 + 1)]  # 多一位
    max_num = 0
    p = 0
    for i in range(lstr1):
        for j in range(lstr2):
            if str1[i] == str2[j]:
                record[i + 1][j + 1] = record[i][j] + 1
                if record[i + 1][j + 1] > max_num:
                    max_num = record[i + 1][j + 1]
                    p = i + 1

    return str1[p - max_num: p], max_num


class DocxTemplateRender(object):
    def __init__(self, filepath: str = None, file_content: IO[bytes] = None):
        self.filepath = filepath
        self.file_content = file_content
        if self.filepath:
            self.doc = Document(self.filepath)
        else:
            self.doc = Document(self.file_content)

    def render(self, template_def):
        doc = self.doc
        # Logics:
        # key_text: "yy{{请输入公司1的完整名称}}xxx""
        # key_runs: [yy{{, 请输入公司, 1, 的完整名称, }}xxx]
        # src_text: ['aaayyy{{', 请输入公司, 1, 的完整名称, }}xxxbbbb]
        # tgt_rext options:
        #  option 1: ['aaa', 北京数据项素，'', 智能科技有限公司, bbbb]
        #  option 2: ['aaa', 北京数据项素智能科技有限公司, '', '', bbbb]
        # allocate replace mapping

        # 采用找key 方式
        for replace_info in template_def:
            k1 = replace_info[0]
            v1 = replace_info[1]
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if k1 in cell.text:
                            table.rows[i].cells[j].text = cell.text.replace(k1, v1)

            for p in doc.paragraphs:
                p.text = p.text.replace(k1, v1)

        return doc
