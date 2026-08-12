#!/usr/bin/env python3
"""Read out and health-check a .docx, for the BiSheng code interpreter.

    import subprocess, sys
    r = subprocess.run([sys.executable, "skills/bisheng-docx/scripts/inspect_docx.py",
                        "output/x.docx"], capture_output=True, text=True)
    print(r.stdout or "(no stdout)")
    print(r.stderr[-2000:] if r.stderr else "(no stderr)")

Flags: ``--content-only`` (dump only), ``--check-only`` (checks only),
``--max-paras N``.

Replaces two things the official skill relies on and this environment lacks:
``pandoc -t markdown`` for reading (pandoc exists in the release image but not
on every hand-built host, and it drops the structural detail needed to plan an
edit), and "render to JPG and look at every page", which assumes a model that
can see images.

The single highest-value check is the missing ``w:eastAsia`` font: it is
invisible in python-docx, invisible in extracted text, and shows up only when
the user opens the file and finds every Chinese character in the wrong face.
The check walks the style inheritance chain, so a run that legitimately
inherits a CJK face from ``Normal`` is not reported.

Always exits 0 — the executor discards stdout on a non-zero exit.
"""

import argparse
import os
import sys
import traceback

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:  # pragma: no cover
    print("[FATAL] python-docx 不可用；无法体检。")
    sys.exit(0)

PLACEHOLDERS = ["待填", "待补", "TODO", "XXX", "占位", "请填写", "TBD", "lorem ipsum", "示例文字"]
CJK_RANGE = ("一", "鿿")
MIN_RUN_PT = 9.0  # design-zh: 图注 9pt 是最小合法字号
MIN_BODY_PT = 10.5  # design-zh: 正文 11pt
REPEAT_HEADER_ROWS = 12  # 超过这个行数的表跨页时必须重复表头

findings: list[tuple[str, str]] = []


def add(level: str, message: str) -> None:
    findings.append((level, message))


def has_cjk(text: str) -> bool:
    return any(CJK_RANGE[0] <= ch <= CJK_RANGE[1] for ch in text)


def _east_asia_of(rpr) -> str | None:
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn("w:eastAsia"))


def run_east_asia(run) -> str | None:
    """The w:eastAsia face written on this run itself, if any."""
    return _east_asia_of(run._element.find(qn("w:rPr")))


def style_east_asia(style) -> str | None:
    """The w:eastAsia face written on this style itself (no inheritance)."""
    if style is None:
        return None
    return _east_asia_of(style.element.find(qn("w:rPr")))


def inherited_east_asia(style) -> str | None:
    """Walk basedOn up the chain: a run under Heading 1 inherits Normal's face."""
    depth = 0
    while style is not None and depth < 8:
        face = style_east_asia(style)
        if face:
            return face
        try:
            style = style.base_style
        except (KeyError, AttributeError):
            return None
        depth += 1
    return None


def style_size_pt(style) -> float | None:
    try:
        size = style.font.size
    except (KeyError, AttributeError):
        return None
    return None if size is None else size.pt


def iter_block_paragraphs(doc):
    """Body paragraphs plus every paragraph inside a table cell."""
    for paragraph in doc.paragraphs:
        yield paragraph, None
    for t_index, table in enumerate(doc.tables):
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    yield paragraph, f"表{t_index + 1}[{r_index + 1},{c_index + 1}]"


def heading_levels(doc) -> list[tuple[int, str]]:
    levels: list[tuple[int, str]] = []
    for paragraph in doc.paragraphs:
        name = paragraph.style.name if paragraph.style is not None else ""
        if name.startswith("Heading ") and paragraph.text.strip():
            try:
                levels.append((int(name.split()[-1]), paragraph.text.strip()))
            except ValueError:
                continue
    return levels


def header_footer_xml(doc) -> str:
    """Raw XML of every header/footer part — PAGE fields do NOT live in document.xml."""
    chunks = []
    for part in doc.part.package.iter_parts():
        name = str(getattr(part, "partname", ""))
        if "header" in name or "footer" in name:
            try:
                chunks.append(part.blob.decode("utf-8", "ignore"))
            except Exception:
                continue
    return "\n".join(chunks)


def field_instructions(doc) -> list[str]:
    return [element.text or "" for element in doc.element.iter(qn("w:instrText"))]


# --------------------------------------------------------------------------- #
# Content read-out
# --------------------------------------------------------------------------- #


def dump_content(path: str, max_paras: int) -> None:
    doc = Document(path)
    section = doc.sections[0]
    width_cm = section.page_width / 360000
    height_cm = section.page_height / 360000

    print("=== 内容 ===")
    print(
        f"{os.path.basename(path)}  页面 {width_cm:.1f}×{height_cm:.1f}cm  "
        f"页边距 上{section.top_margin / 360000:.1f} 下{section.bottom_margin / 360000:.1f} "
        f"左{section.left_margin / 360000:.1f} 右{section.right_margin / 360000:.1f}cm  "
        f"段落 {len(doc.paragraphs)}  表格 {len(doc.tables)}  图片 {len(doc.inline_shapes)}"
    )
    print()

    shown = 0
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if shown >= max_paras:
            print("   … 其余段落未显示（--max-paras 可调）")
            break
        style = paragraph.style.name if paragraph.style is not None else "?"
        tag = f"[{style}]" if style != "Normal" else ""
        print(f"P{index:<4} {tag} {text if len(text) <= 110 else text[:107] + '…'}")
        shown += 1

    for t_index, table in enumerate(doc.tables):
        print(f"\n## 表 {t_index + 1} · {len(table.rows)} 行 × {len(table.columns)} 列")
        for row in table.rows[:8]:
            cells = [c.text.strip().replace("\n", " ")[:24] for c in row.cells]
            print("   | " + " | ".join(cells) + " |")
        if len(table.rows) > 8:
            print(f"   … 其余 {len(table.rows) - 8} 行未显示")


# --------------------------------------------------------------------------- #
# Health checks
# --------------------------------------------------------------------------- #


def check_fonts(doc) -> None:
    """The headline check: every Chinese run must end up with a w:eastAsia face."""
    if not style_east_asia(doc.styles["Normal"]):
        add(
            "WARN",
            "Normal 样式没有设置 w:eastAsia 中文字体 —— 中文会退回主题字体（通常是等线/宋体），"
            "和你指定的西文字体对不上。在构建脚本里调 apply_chinese_defaults(doc) 一次性设好。",
        )

    uncovered: list[str] = []  # 完全没有中文字体可继承
    mismatched: dict[str, str] = {}  # run 写了西文字体、中文却继承自样式
    style_gaps: dict[str, str] = {}  # 该段落样式整条继承链都没有中文字体
    checked = 0
    for paragraph, where in iter_block_paragraphs(doc):
        style = paragraph.style
        style_face = inherited_east_asia(style)
        for run in paragraph.runs:
            if not has_cjk(run.text):
                continue
            checked += 1
            label = where or f"「{run.text.strip()[:16]}」"
            east = run_east_asia(run)
            latin = run.font.name
            if east:
                continue
            if not style_face:
                if latin and label not in uncovered:
                    uncovered.append(label)
                elif not latin:
                    style_gaps.setdefault(style.name if style is not None else "?", label)
            elif latin and latin != style_face:
                mismatched.setdefault(label, f"{latin}→实际 {style_face}")

    if uncovered:
        add(
            "ERROR",
            f"{len(uncovered)} 处中文 run 只设了西文字体、没设 w:eastAsia，且样式里也没有中文字体："
            f"{', '.join(uncovered[:6])}{' …' if len(uncovered) > 6 else ''}。"
            "`run.font.name = '微软雅黑'` 对中文无效 —— 改用 set_run_font(run, '微软雅黑', size_pt=11)。",
        )
    if style_gaps:
        names = ", ".join(f"{k}（如{v}）" for k, v in list(style_gaps.items())[:4])
        add(
            "ERROR",
            f"{len(style_gaps)} 个段落样式的继承链里没有 w:eastAsia 中文字体：{names}。"
            "这些中文会退回主题字体，用户打开就是「字体乱了」。"
            "调 apply_chinese_defaults(doc)（它会把 Normal 和 Heading 1–4 一次设好），"
            "或用 add_heading_cn / add_body 写内容。",
        )
    if mismatched:
        items = ", ".join(f"{k}:{v}" for k, v in list(mismatched.items())[:4])
        add(
            "WARN",
            f"{len(mismatched)} 处中文 run 写的是西文字体、中文实际用的是样式继承来的另一种字体：{items}。"
            "同一段中西文字体不一致会看出明显断层 —— 用 set_run_font() 把三个字段一起设。",
        )
    if checked == 0:
        add("INFO", "文档里没有中文 run（纯英文文档？）")

    small: list[str] = []
    for paragraph, where in iter_block_paragraphs(doc):
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            size = run.font.size
            if size is not None and size.pt < MIN_RUN_PT:
                small.append(f"{where or run.text.strip()[:12]}={size.pt:g}pt")
    if small:
        add(
            "WARN",
            f"{len(small)} 处文字字号小于 {MIN_RUN_PT:g}pt：{', '.join(small[:6])}"
            f"{' …' if len(small) > 6 else ''}。打印后基本读不清 —— "
            f"正文 11pt、表格 10pt、图注 9pt，最小不要低于 {MIN_RUN_PT:g}pt。",
        )
    normal_pt = style_size_pt(doc.styles["Normal"])
    if normal_pt is not None and normal_pt < MIN_BODY_PT:
        add(
            "WARN",
            f"Normal 样式字号 {normal_pt:g}pt 偏小（正文规范 11pt）。"
            "在 apply_chinese_defaults(doc, body_pt=11) 里调回来。",
        )


def check_headings(doc) -> None:
    levels = heading_levels(doc)
    if not levels:
        add(
            "WARN",
            "文档里没有用内置 Heading 样式的标题 —— 目录生成不出来，导航窗格也是空的。"
            "用 add_heading_cn(doc, '一、xxx', 2)，不要用加粗的普通段落冒充标题。",
        )
        return
    previous = 0
    for level, text in levels:
        if previous and level > previous + 1:
            add(
                "WARN",
                f"标题层级从 H{previous} 直接跳到 H{level}：「{text[:24]}」—— 中间补一级，或把它降到 H{previous + 1}",
            )
        previous = level


def _cell_shaded(cell) -> bool:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return False
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return False
    fill = (shd.get(qn("w:fill")) or "auto").lower()
    return fill not in ("auto", "ffffff", "")


def _row_has_bold(row) -> bool:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.font.bold or (run.style is not None and getattr(run.style.font, "bold", None)):
                    return True
    return False


def _row_repeats_as_header(row) -> bool:
    tr_pr = row._tr.find(qn("w:trPr"))
    return tr_pr is not None and tr_pr.find(qn("w:tblHeader")) is not None


def check_tables(doc, section) -> None:
    usable_cm = (section.page_width - section.left_margin - section.right_margin) / 360000
    for index, table in enumerate(doc.tables):
        if not table.rows:
            add("WARN", f"表 {index + 1} 是空的 —— 删掉它，或把数据填进去")
            continue
        widths = [c.width for c in table.rows[0].cells]
        if all(w is not None for w in widths):
            total_cm = sum(w / 360000 for w in widths)
            if total_cm > usable_cm + 0.2:
                add(
                    "ERROR",
                    f"表 {index + 1} 总宽 {total_cm:.1f}cm 超过版心 {usable_cm:.1f}cm → 右侧列会被截出页面。"
                    f"把 widths_cm 的合计压到 {usable_cm:.1f}cm 以内（用 content_width_cm(section) 取准确值）。",
                )
        elif any(w is None for w in widths):
            add(
                "INFO",
                f"表 {index + 1} 没有设列宽，Word 会自动分配；列多时中文列容易被压成竖排。"
                f"给 add_table 传 widths_cm=[…]（合计 ≤ {usable_cm:.1f}cm）。",
            )
        header = table.rows[0]
        if not _cell_shaded(header.cells[0]) and not _row_has_bold(header):
            add(
                "WARN",
                f"表 {index + 1} 首行不像表头（既没有加粗也没有底纹）—— 读者分不清表头和数据。"
                "用 add_table(doc, headers=[…], rows=[…])，它会自动加深底白字加粗。",
            )
        if len(table.rows) > REPEAT_HEADER_ROWS and not _row_repeats_as_header(header):
            add(
                "INFO",
                f"表 {index + 1} 有 {len(table.rows)} 行，跨页后第二页没有表头。"
                "add_table 已自动设置重复表头；手写表格用 repeat_header_row(table) 补上。",
            )
        for r_index, row in enumerate(table.rows):
            if len(row.cells) != len(table.columns):
                add(
                    "WARN",
                    f"表 {index + 1} 第 {r_index + 1} 行单元格数 {len(row.cells)} 与列数 {len(table.columns)} 不一致"
                    "（可能有合并单元格）—— 确认是有意为之，否则补齐这一行。",
                )
                break


def check_images(doc, section) -> None:
    usable_cm = (section.page_width - section.left_margin - section.right_margin) / 360000
    for index, shape in enumerate(doc.inline_shapes):
        width_cm = shape.width / 360000
        if width_cm > usable_cm + 0.2:
            add(
                "ERROR",
                f"图 {index + 1} 宽 {width_cm:.1f}cm 超过版心 {usable_cm:.1f}cm → 会溢出页面。"
                "用 add_image_fitted(doc, path, section=section) 按版心等比缩放。",
            )


def check_text(doc) -> None:
    empty_streak = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            empty_streak += 1
            if empty_streak == 3:
                add("INFO", "有连续 3 个以上空段落 —— 改用段前段后间距（space_before/space_after）分隔内容")
        else:
            empty_streak = 0

    for paragraph, where in iter_block_paragraphs(doc):
        text = paragraph.text.strip()
        if not text:
            continue
        prefix = f"{where} " if where else ""
        for token in PLACEHOLDERS:
            if token in text:
                add("WARN", f"{prefix}残留占位文本「{token}」：{text[:40]} —— 换成真实内容，或整段删掉")
                break
        if "\n" in paragraph.text:
            add(
                "WARN", f"{prefix}段落里有换行符而不是独立段落：{text[:40]} —— 每段一个 Paragraph，或用 run.add_break()"
            )


def check_fields(doc) -> None:
    instructions = field_instructions(doc)
    hdr_ftr = header_footer_xml(doc)
    has_toc = any("TOC" in text for text in instructions)
    if has_toc:
        if "updateFields" not in doc.settings.element.xml:
            add(
                "ERROR",
                "有目录域但没开 updateFields —— 用户打开时目录仍是占位文字，看起来像坏了。"
                "调用 enable_update_fields(doc)（add_toc 已内置）。",
            )
        if not heading_levels(doc):
            add(
                "ERROR",
                "有目录域但全文没有内置 Heading 样式的标题 → 目录会是空的。"
                "标题一律用 add_heading_cn(doc, text, level)。",
            )
    has_page_field = any("PAGE" in text for text in instructions) or "PAGE" in hdr_ftr
    if not has_page_field and len(doc.paragraphs) > 40:
        add(
            "INFO",
            f"文档有 {len(doc.paragraphs)} 个段落但页脚没有页码域 —— 调 add_page_number_footer(section)。",
        )


def run_checks(path: str) -> None:
    doc = Document(path)
    section = doc.sections[0]
    check_fonts(doc)
    check_headings(doc)
    check_tables(doc, section)
    check_images(doc, section)
    check_text(doc)
    check_fields(doc)

    print()
    print("=== 体检 ===")

    order = {"ERROR": 0, "WARN": 1, "INFO": 2}
    counts = {"ERROR": 0, "WARN": 0, "INFO": 0}
    for level, _ in findings:
        counts[level] += 1
    for level, message in sorted(findings, key=lambda f: order[f[0]]):
        print(f"[{level}] {message}")
    if not findings:
        print("（无发现）")

    print()
    print(f"合计: {counts['ERROR']} ERROR / {counts['WARN']} WARN / {counts['INFO']} INFO")
    if counts["ERROR"]:
        print(f"结论: 不通过 —— {counts['ERROR']} 项 ERROR 必须修完。改构建脚本重新生成，再跑一次本脚本。")
    else:
        print("结论: 通过 —— 无必须修复项。WARN 逐条复核后即可交付。")


def main() -> None:
    parser = argparse.ArgumentParser(description="读出并体检一个 .docx")
    parser.add_argument("path")
    parser.add_argument("--content-only", action="store_true")
    parser.add_argument("--check-only", action="store_true")
    parser.add_argument("--max-paras", type=int, default=80)
    args = parser.parse_args()

    if not os.path.exists(args.path):
        print(f"[FATAL] 文件不存在: {args.path}")
        print("提示：代码执行器 cwd 就是工作区根，用相对路径 `output/x.docx`，不要写 `/output/x.docx`。")
        return

    try:
        if not args.check_only:
            dump_content(args.path, args.max_paras)
        if not args.content_only:
            run_checks(args.path)
    except Exception:
        print("[FATAL] 体检过程本身出错：")
        traceback.print_exc(file=sys.stdout)


if __name__ == "__main__":
    main()
    sys.exit(0)
